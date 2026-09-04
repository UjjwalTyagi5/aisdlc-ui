"""Every chat agent that accepts attachments must actually read them.

The Plan route was typed `user_message_with_files` and read everything on the
message EXCEPT the files, so an attached BRD appeared as a chip in the transcript
while the agent replied "I don't see a BRD attached to your message" — asking for
the document the user had just given it. Requirements and Design each carried their
own near-identical copy of the extraction block; Plan was written without one and
nothing failed loudly.
"""
import pathlib
import re

import pytest
from docx import Document

from shared.tools.document_tools import (
    attachment_message_contents,
    attachment_paths_from_context,
)


@pytest.fixture
def brd(tmp_path) -> str:
    doc = Document()
    doc.add_heading("QuickLink URL Shortener — BRD", 0)
    doc.add_paragraph("REQ-01: Users can shorten a long URL into a 7-character slug.")
    path = tmp_path / "QuickLink_BRD.docx"
    doc.save(str(path))
    return str(path)


def test_a_docx_becomes_content_the_model_can_use(brd):
    contents = attachment_message_contents([brd])

    assert len(contents) == 1
    assert "use their content directly" in contents[0]
    assert "QuickLink_BRD.docx" in contents[0]
    assert "REQ-01" in contents[0]  # the actual text, not just the name


def test_nothing_attached_produces_no_message():
    """The caller appends nothing rather than an empty instruction."""
    assert attachment_message_contents([]) == []
    assert attachment_message_contents(None) == []


def test_an_unreadable_file_is_reported_instead_of_faked(tmp_path):
    """Extraction returns a readable PLACEHOLDER on failure, so a truthiness check
    would announce a screenshot to the agent as document content."""
    png = tmp_path / "shot.png"
    png.write_bytes(b"\x89PNG\r\n\x1a\n" + b"0" * 40)

    contents = attachment_message_contents([str(png)])

    assert len(contents) == 1
    assert "could not be read as text" in contents[0]
    assert "do not claim to have looked at it" in contents[0]


def test_a_readable_and_an_unreadable_file_are_reported_separately(brd, tmp_path):
    png = tmp_path / "shot.png"
    png.write_bytes(b"\x89PNG\r\n\x1a\n" + b"0" * 40)

    contents = attachment_message_contents([brd, str(png)])

    assert len(contents) == 2
    assert "REQ-01" in contents[0]
    assert "shot.png" in contents[1]


def test_a_missing_path_does_not_raise(tmp_path):
    contents = attachment_message_contents([str(tmp_path / "gone.docx")])

    assert len(contents) == 1
    assert "could not be read" in contents[0]


def test_extracted_text_is_capped(tmp_path):
    """Several attachments must not be able to exhaust a context window."""
    big = tmp_path / "big.txt"
    big.write_text("x" * 100_000, encoding="utf-8")

    contents = attachment_message_contents([str(big)])

    assert len(contents[0]) < 25_000


def test_attachment_paths_are_read_off_the_pipeline_context():
    ctx = {"attachments": [{"path": "/a/b.docx"}, {"name": "no path"}, "junk"]}

    assert attachment_paths_from_context(ctx) == ["/a/b.docx"]
    assert attachment_paths_from_context(None) == []
    assert attachment_paths_from_context({"attachments": None}) == []


# The WEBSOCKET turn handler of each chat agent — the path the chat drawer actually
# uses. Named individually because "the module mentions attachments somewhere" is not
# the property that matters, and asserting only that is how this bug survived a fix:
# the Plan agent got the block in its REST handler while the socket the UI talks to
# kept dropping every file.
_WS_TURN_HANDLERS = [
    ("agents_orchestrator.pm_agent.pm_agent_api", "_process_turn_ws"),
    ("agents_orchestrator.requirements_agent.requirements_agent_api",
     "_process_user_message_ws"),
    ("agents_orchestrator.design_architecture_agent.design_architecture_agent_api",
     "_process_user_message_ws"),
]


@pytest.mark.parametrize("module_name,handler_name", _WS_TURN_HANDLERS)
def test_every_websocket_turn_handler_reads_its_attachments(module_name, handler_name):
    import importlib
    import inspect

    module = importlib.import_module(module_name)
    handler = getattr(module, handler_name, None)
    assert handler is not None, f"{module_name} has no {handler_name}"

    source = inspect.getsource(handler)
    assert "attachment_message_contents(" in source, (
        f"{module_name}.{handler_name} does not read attachments — a file attached "
        f"in the chat drawer will be silently dropped"
    )
