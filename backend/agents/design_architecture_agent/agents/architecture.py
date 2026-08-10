import os
import time
import asyncio
import aiohttp
import aiofiles
import requests
import json
import re
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_COLOR_INDEX
from io import BytesIO
from xhtml2pdf import pisa
import textwrap
import docx
import markdown
from dotenv import load_dotenv
from typing import TypedDict, Annotated, List, Dict, Any
import google.generativeai as genai
from google.ai.generativelanguage import File
from langgraph.checkpoint.memory import MemorySaver
from langchain_core.messages import BaseMessage, HumanMessage, ToolMessage, AIMessage
from langchain_core.tools import tool
from langchain_google_genai import ChatGoogleGenerativeAI
from langgraph.graph import StateGraph, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode
from langgraph.prebuilt import InjectedState
from langgraph.types import Command
from config.ws_helper import set_session_id, broadcast_log
from config.connection_manager import manager
from agents.design_architecture_agent.prompts.architecture_generation import ARCH_GEN_PROMPT
from docxtpl import DocxTemplate, InlineImage, RichText
import pickle
from agents.design_architecture_agent.config import shared
from docx2pdf import convert
import mermaid as mmd 
from mermaid.graph import Graph
import mermaid.flowchart as fl
from config.ws_helper import set_session_id, broadcast_log, get_user_id, get_session_id
import certifi
import inspect
currentdir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
parentdir = os.path.dirname(currentdir)
superparentdir = os.path.dirname(parentdir)
supersuperparentdir = os.path.dirname(superparentdir)
from config.env import AGENTIC_BASE_URL
from config.agent_context import get_agent_folder, set_agent_folder
from config import sdlcSettings
from config.env import GOOGLE_API_KEY_DESIGN
esett = sdlcSettings()

genai.configure(api_key=GOOGLE_API_KEY_DESIGN)
config = genai.types.GenerationConfig(temperature=.25)
model = genai.GenerativeModel(model_name="gemini-2.5-flash")

# Enhanced file_generated broadcast with file size
async def broadcast_file_generated(session_id: str, filename: str, file_path: str):
    """Enhanced function to broadcast file generation with file size"""
    try:
        user_id = get_user_id()
        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
        file_url = f"{AGENTIC_BASE_URL}/generated/{user_id}/{get_agent_folder()}/{session_id}/output/{filename}"
        await manager.broadcast({
            "type": "file_generated",
            "session_id": session_id,
            "filename": filename,
            "url": file_url,
            "file_size": file_size,
            "message": f"Generated file: {filename}"
        })
        
        print(f"DEBUG: Broadcasted file generation: {filename} ({file_size} bytes)")
        
    except Exception as e:
        print(f"ERROR: Failed to broadcast file generation: {str(e)}")


def replace_parentheses_with_quotes(text):
    """
    Replaces all '(' and ')' characters in the input text with double quotes.

    Args:
        text (str): The input string.

    Returns:
        str: Modified string with '(' and ')' replaced by '"'.
    """
    return text.replace("(", '').replace(")", '').replace("<br>","\n").replace("</br>","")

def fix_mermaid_syntax(mermaid_code):
    fixed_lines = []
    node_pattern = re.compile(r'([a-zA-Z0-9_]+)\s+([a-zA-Z0-9\s&]+)')  # matches: ID Label

    for line in mermaid_code.splitlines():
        # Skip empty lines or subgraph declarations
        if not line.strip() or line.strip().startswith("subgraph") or "--" not in line:
            fixed_lines.append(line)
            continue

        # Fix node names with spaces not wrapped in brackets
        matches = re.findall(r'([a-zA-Z0-9_]+)\s+([a-zA-Z][a-zA-Z0-9\s&]+)', line)
        for match in matches:
            node_id, label = match
            full_text = f'{node_id} {label}'
            replacement = f'{node_id}["{label}"]'
            line = line.replace(full_text, replacement)

        # Optionally fix double node names like "CustomersCustomers"
        line = re.sub(r'\b(\w+)\1\b', r'\1', line)

        fixed_lines.append(line)

    return "\n".join(fixed_lines)

# Async version of mermaid_to_png
async def mermaid_to_png_async(graph: str,session_id):
    """Async version of mermaid to PNG conversion"""
    user_id = get_user_id()
    # graph = replace_parentheses_with_quotes(graph)
    # graph = fix_mermaid_syntax(graph)
    graph ="""
      graph TD
    subgraph External Systems
        OMS[SwiftServe Order Management System]
        TrafficAPI[External Traffic APIs]
        WeatherAPI[External Weather Services]
    end

    subgraph SwiftRoutePlatform_CloudNative[Swift Route Platform_Cloud_Native]
        LoadBalancer[Cloud Load Balancer]
        Gateway[API Gateway]

        subgraph CoreServices
            DataIngestion[Data Ingestion & Processing]
            ROE[Intelligent Route Optimization Engine]
            FMS[Fleet Management Service]
            NotificationSvc[Notification Service]
            IntegrationSvc[OMS Integration Service]
        end

        Database[PostgreSQL Database]
        Cache[Redis Cache]
        Queue[Message Queue e.g. Kafka SQS]
        ObjectStorage[Object Storage e.g. S3]
    end

    subgraph User Interfaces
        DispatcherUI[Dispatcher Web Interface]
        DriverApp[Driver Mobile Application]
        CustomerPortal[Customer Web Portal]
    end

    OMS -- Pull Orders --> IntegrationSvc
    IntegrationSvc -- Push Status Updates --> OMS

    TrafficAPI -- Ingest Data --> DataIngestion
    WeatherAPI -- Ingest Data --> DataIngestion
    DataIngestion -- Feeds Real-time Data --> ROE
    DataIngestion -- Stores Data --> Database

    ROE -- Optimizes Routes --> FMS
    FMS -- Assigns Drivers, Manages Fleet --> Database
    FMS -- Real-time Updates --> Queue

    Gateway -- API Endpoints --> CoreServices
    DispatcherUI -- Manages Fleet & Routes --> Gateway
    DriverApp -- Receives Routes, Sends Status --> Gateway
    CustomerPortal -- Live Tracking & Updates --> Gateway

    CoreServices -- Stores Persistent Data --> Database
    CoreServices -- Caches Frequent Data --> Cache
    CoreServices -- Async Processing/Notifications --> Queue

    DriverApp -- Location Data, PoD --> ObjectStorage
    NotificationSvc -- Sends Alerts/Updates --> CustomerPortal"""
    
    print(graph,"@@##")
    print(session_id,"!!@@###!!!W@@")
    def _mermaid_convert():
        old_get = requests.get
        def new_get(*args, **kwargs):
            kwargs['verify'] = False
            return old_get(*args, **kwargs)
        requests.get = new_get

        mermaid = mmd.Mermaid(graph)
        out_path = f"{esett.FILES}/{user_id}/{get_agent_folder()}/{session_id}/output/"
        # out_path = f"C:/Users/mprasad053/Downloads/outputs/{user_id}/{get_agent_folder()}/{session_id}/output"
        if not os.path.exists(out_path):
            os.makedirs(out_path)
        # output_path = f"C:/Users/mprasad053/Downloads/outputs/{user_id}/{get_agent_folder()}/{session_id}/output/mermaid.png"
        output_path = f"{esett.FILES}/{user_id}/{get_agent_folder()}/{session_id}/output/mermaid.png"
        print(output_path,"@@##$$%%%%%%%")
        mermaid.to_png(output_path)
        filename = os.path.basename(output_path)
        return output_path
    
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _mermaid_convert)

def convert_to_pdf(input_file_path: str, output_directory: str) -> str:
    """
    Convert a DOCX file to PDF using docx2pdf.
    """
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    base_name, _ = os.path.splitext(os.path.basename(input_file_path))
    output_file_path = os.path.join(output_directory, f"{base_name}.pdf")

    try:
        convert(input_file_path, output_directory)
        return output_file_path
    except Exception as e:
        raise RuntimeError(f"Conversion failed: {e}")

def add_hyperlink(paragraph, text, url):
    """Adds a hyperlink to a paragraph."""
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
   
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
   
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    run.append(rPr)
    run.append(OxmlElement('w:t', text=text))
   
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink

async def markdown_to_docx_async(markdown_string: str, docx_path: str):
    """
    Async version of markdown to DOCX conversion with near-full feature support.
    """
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn

    print(f"Starting async markdown to docx conversion for: {docx_path}")
    dedented_markdown = textwrap.dedent(markdown_string).strip()
    doc = Document()
    html = markdown.markdown(dedented_markdown, extensions=['extra', 'sane_lists', 'tables', 'fenced_code', 'codehilite'])
    soup = BeautifulSoup(html, 'html.parser')

    def _add_inline_content(paragraph, element: Tag):
        """Recursively adds formatted runs to a paragraph."""
        for child in element.children:
            if isinstance(child, NavigableString):
                if child.strip():
                    run = paragraph.add_run(child)
                    curr = child.parent
                    while curr is not None and curr.name != element.name:
                        if curr.name in ['strong', 'b']: run.bold = True
                        if curr.name in ['em', 'i']: run.italic = True
                        if curr.name in ['s', 'del']: run.strike = True
                        if curr.name == 'code':
                            run.font.name = 'Courier New'
                        curr = curr.parent
            elif isinstance(child, Tag):
                if child.name == 'br':
                    paragraph.add_run().add_break()
                elif child.name == 'a':
                    add_hyperlink(paragraph, child.get_text(), child.get('href'))
                else:
                    _add_inline_content(paragraph, child)

    async def _parse_element(element, doc, list_info=None):
        """Async version of element parsing."""
        if not hasattr(element, 'name') or element.name is None:
            return

        if element.name in [f'h{i}' for i in range(1, 7)]:
            level = int(element.name[1])
            p = doc.add_heading(level=level)
            _add_inline_content(p, element)
            print(f"Added heading level {level}")
            await asyncio.sleep(0.001)
        elif element.name == 'p':
            p = doc.add_paragraph()
            _add_inline_content(p, element)
            print("Added paragraph")
            await asyncio.sleep(0.001)
        elif element.name in ['ul', 'ol']:
            list_style = 'List Bullet' if element.name == 'ul' else 'List Number'
            parent_level = list_info[1] if list_info else -1
            for li in element.find_all('li', recursive=False):
                await _parse_element(li, doc, list_info=(list_style, parent_level + 1))
        elif element.name == 'li':
            if list_info:
                style, level = list_info
                p = doc.add_paragraph(style=style)
                p.paragraph_format.left_indent = Inches(0.5 * level)
                _add_inline_content(p, element)
                for nested_list in element.find_all(['ul', 'ol'], recursive=False):
                    await _parse_element(nested_list, doc, list_info)
                print(f"Added list item at level {level}")
                await asyncio.sleep(0.001)
        elif element.name == 'table':
            rows_data = element.find_all('tr')
            if not rows_data: return
            num_cols = len(rows_data[0].find_all(['th', 'td']))
            table = doc.add_table(rows=len(rows_data), cols=num_cols)
            table.style = 'Table Grid'
            for i, row_element in enumerate(rows_data):
                cells = row_element.find_all(['th', 'td'])
                for j, cell_element in enumerate(cells):
                    p = table.cell(i, j).paragraphs[0]
                    _add_inline_content(p, cell_element)
            print(f"Added table with {len(rows_data)} rows and {num_cols} columns")
            await asyncio.sleep(0.001)
        elif element.name == 'img':
            src = element.get('src')
            if src:
                try:
                    async with aiohttp.ClientSession() as session:
                        async with session.get(src) as response:
                            if response.status == 200:
                                image_data = await response.read()
                                doc.add_picture(BytesIO(image_data), width=Inches(5.0))
                                print(f"Added image from {src}")
                            else:
                                print(f"Warning: Could not fetch image from {src}. Status: {response.status}")
                                doc.add_paragraph(f"[Image not found: {src}]").add_run().italic = True
                except Exception as e:
                    print(f"Warning: Could not fetch image from {src}. Error: {e}")
                    doc.add_paragraph(f"[Image not found: {src}]").add_run().italic = True
        elif element.name == 'hr':
            doc.add_page_break()
            print("Added page break")
            await asyncio.sleep(0.001)
        elif element.name == 'blockquote':
            p = doc.add_paragraph(style='Quote')
            p.paragraph_format.left_indent = Inches(0.5)
            for child in element.children:
                await _parse_element(child, doc)
            print("Added blockquote")
            await asyncio.sleep(0.001)
        elif element.name == 'pre':
            code_text = element.get_text()
            p = doc.add_paragraph(style='No Spacing')
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            print("Added code block")
            await asyncio.sleep(0.001)

    # Process all top-level tags from the parsed HTML
    for element in soup.contents:
        await _parse_element(element, doc)
   
    try:
        session_id = get_session_id()
        print(f"Saving document to {docx_path}")
        doc.save(docx_path)
        filename = os.path.basename(docx_path)
        await broadcast_file_generated(session_id, filename, docx_path)
        print(f"Successfully saved document to {docx_path}")
    except Exception as e:
        print(f"An exception occurred while saving the doc file: {e}")
    return f"Successfully converted Markdown to '{docx_path}'"

class AgentState(TypedDict):
    messages: Annotated[List[BaseMessage], add_messages]

@tool
async def upload_file(local_path: str):
    """
    Uploads a local file to the cloud service so it can be used by other tools
    It stores the file handle in the agent's state for later use

    Args:
        local_path (str): The local path of the file to upload

    Returns:
        str: The file handle name that can be used by other tools
    """
    user_id = get_user_id()
    broadcast_log(manager, f"---Executing tool: Upload_file for path: {local_path}---", level="INFO")

    if not os.path.exists(local_path):
        error_msg = f"Error: local file not found at {local_path}"
        print(error_msg)
        return error_msg

    try:
        print(f"Starting upload for: {local_path}")
        loop = asyncio.get_event_loop()
        file = await loop.run_in_executor(None, genai.upload_file, local_path)
        print("File uploaded, waiting for processing...")
       
        while file.state.name == "PROCESSING":
            print("Still processing...")
            await asyncio.sleep(2)
            file = await loop.run_in_executor(None, genai.get_file, file.name)

        if file.state.name == "FAILED":
            error_msg = f"Error: File upload has failed for {local_path}"
            print(error_msg)
            return error_msg

        print(f"File {file.name} is now ACTIVE")
        user_id = get_user_id()
        session_id = get_session_id()
        file_path = f"{esett.FILES}/{user_id}/{get_agent_folder()}/{session_id}/input"
        file_path_file = f"{esett.FILES}/{user_id}/{get_agent_folder()}/{session_id}/input/files"

        # Save the file handle
        if os.path.exists(file_path):
            print(f"The folder '{file_path}' already exists.")
        else:
            os.makedirs(file_path)
        if os.path.exists(file_path_file):
            print(f"The folder '{file_path_file}' already exists.")
        else:
            os.makedirs(file_path_file)

        final_path = os.path.join(file_path, file.name)
        
        async with aiofiles.open(final_path, "wb") as f:
            await f.write(pickle.dumps(file))
            print(f"File handle saved: {final_path}")

        return final_path

    except Exception as e:
        error_msg = f"Exception during upload: {str(e)}"
        print(error_msg)
        return error_msg

@tool
async def delete_file(file_name: str):
    """
    Deletes a file from the cloud service using its file name, all uploaded files must be deleted after their use is done
   
    Args:
        file_name (str): The name of the file to delete, use the name returned by upload file
    """    
    print(f"---Executing tool: delete file for name: {file_name}---")
    parts = file_name.split("files")

    if len(parts) > 2:
        # Rebuild the filename from after the second 'files'
        # parts[0] is before 1st 'files'
        # parts[1] is between 1st and 2nd 'files'
        # parts[2] onwards is after 2nd 'files'
        new_file_name = "files".join(parts[2:])  # re-attach in case 'files' is in path filenames/etc.
        # Strip any leading separators or whitespace
        new_file_name = new_file_name.lstrip("/\\ ")
        print(f"Extracted filename after second 'files': {new_file_name}")
        file_name = new_file_name
    elif len(parts) == 2:
        # If there are exactly 2 parts, take substring after first 'files' (optional)
        new_file_name = parts[1].lstrip("/\\ ")
        print(f"Extracted filename after first 'files': {new_file_name}")
        file_name = new_file_name
    else:
        # filename doesn't contain 'files' or contains only once, use as is
        print("File name contains less than two occurrences of 'files', using as is.")\

    # The local file path is the full path argument with forward slashes fixed
    file_path = file_name.replace("\\", "/")

    # Check local file exists before deleting locally
    if not os.path.exists(file_path):
        print(f"Warning: local file path '{file_path}' not found. Skipping local delete.")
    else:
        # Delete local file - async executor for blocking call
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, os.remove, file_path)
        print(f"Local file '{file_path}' deleted.")

    # Call genai.delete_file with file ID (should be short id, no full path)
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, genai.delete_file, file_name)
    print(f"GenAI file with ID '{file_name}' deleted.")

    return f"Successfully deleted file {file_path} and remote file {file_name}"

@tool
async def generate_architechture(file_names: List[str], custom_prompt: str):
    """
    Generates project architecture by analyzing a collection of files,
    which can include transcripts, audios and meeting recordings and pdfs.
   
    Args:
        file_names (List[str]): A list of file names to analyze
                                These names must be obtained from the upload_file tool
        custom_prompt (str): Additional instructions needed based on users query, can be empty
    """
    broadcast_log(manager, f"---Executing tool: generate_architecture for files: {file_names}---", level="INFO")

    gemini_files = []
    for name in file_names:
        if name.startswith("Error:"):
            print(f"Skipping error message: {name}")
            continue
            
        if os.path.exists(name):
            try:
                async with aiofiles.open(name, "rb") as f:
                    file_data = await f.read()
                    file = pickle.loads(file_data)
                    gemini_files.append(file)
                print(f"Successfully loaded file: {name}")
            except Exception as e:
                print(f"Error loading file {name}: {str(e)}")
                continue
        else:
            return f"Error: file {name} has not yet been uploaded"
            
    if not gemini_files:
        return f"Error: Must provide at least one valid file to generate Architecture"

    global model
    try:
        print("Preparing prompt")
        prompt = ARCH_GEN_PROMPT.format(custom_prompt=custom_prompt)
    except Exception as e:
        print(f"Error formatting prompt: {e}")
        return f"Error formatting prompt: {e}"

    print("File was loaded, generating architecture...")
    try:
        broadcast_log(manager, "Generating architecture content...", level="INFO")
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(None, model.generate_content, [prompt] + gemini_files)
        print(response,"RESPONSE******************************8")
        broadcast_log(manager, "Architecture generation completed", level="INFO")
        
        s = response.text
        async with aiofiles.open("architecture_json.txt", "w", encoding="utf-8") as f:
            await f.write(s)
        
        print("Architecture generated successfully")
        pattern = r"```(.*?)```"
        match = re.search(pattern, s, re.DOTALL)
        new_s = ""
        if match:
            code_block = match.group(1)
            print("Found mermaid diagram")
            shared.mermaid = code_block.replace("mermaid", "")
            replacement = "{{mermaid}}"
            new_s = s.replace(match.group(0), replacement, 1)
        else:
            new_s = s
            
        return new_s
       
    except Exception as e:
        error_msg = f"An exception occurred while generating the Architecture: {str(e)}"
        broadcast_log(manager, error_msg, level="ERROR")
        return error_msg

@tool
async def update_response(file_names: List[str], query: str, content: str):
    """
    A function that can be used to update the content provided based on context, returns the updated content
   
    Args:
    file_names (List[str]): A list of file names obtained from upload file that is provided as context
    query (str): the actual query or the updates the user needs, combine all the needs previously mentioned
    content (str): the actual content that needs to be updated
    """
    print(f"Executing tool: update_response for files: {file_names}---")
    broadcast_log(manager, "Updating response based on context...", level="INFO")
    
    gemini_files = []
    for name in file_names:
        if os.path.exists(name):
            async with aiofiles.open(name, "rb") as f:
                file_data = await f.read()
                file = pickle.loads(file_data)
                gemini_files.append(file)
        else:
            return f"Error: file {name} has not yet been uploaded"

    global model
    prompt = """Given the following content and attached files as context update the content provided to you based on the query
Query: {query}

````````
Content:
{content}
`````````
    """
    prompt = prompt.format(content=content, query=query)
    
    try:
        broadcast_log(manager, "Updating response...", level="INFO")
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(None, model.generate_content, [prompt] + gemini_files)
        broadcast_log(manager, "Response update completed", level="INFO")
        return response.text
       
    except Exception as e:
        error_msg = f"An exception occurred while updating the query: {str(e)}"
        broadcast_log(manager, error_msg, level="ERROR")
        return error_msg

@tool
async def markdowntodoc(content: str, output_path: str):
    """
    Function to convert a HTML or Markdown string into a docx file and save it
    Args:
    content (str): the actual markdown string
    output_path (str): the local output path to store the file: typically file.docx
    """
    return await markdown_to_docx_async(content, output_path)

def clean_and_parse_json(json_string: str) -> dict:
    """
    Attempts to clean a possibly malformed or wrapped JSON string and parse it into a Python dictionary.
    """
    original = json_string

    # Step 1: Strip BOM and clean control characters
    cleaned = json_string.strip().lstrip('\ufeff')
    cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', cleaned)

    # Step 2: Remove Markdown-style code block markers
    cleaned = re.sub(r'^```(?:json|python)?\s*', '', cleaned, flags=re.IGNORECASE | re.MULTILINE)
    cleaned = re.sub(r'\s*```$', '', cleaned, flags=re.MULTILINE).strip()

    # Step 3: Use a streaming JSON decoder to find the first valid JSON
    decoder = json.JSONDecoder()
    for match in re.finditer(r'[\[{]', cleaned):
        try:
            candidate = cleaned[match.start():]
            obj, idx = decoder.raw_decode(candidate)
            if isinstance(obj, dict) or isinstance(obj, list):
                return obj
        except json.JSONDecodeError:
            continue

    # Step 4: Fallback to last-ditch balance-based trimming
    brace_pos = cleaned.find('{')
    bracket_pos = cleaned.find('[')

    if brace_pos == -1 and bracket_pos == -1:
        raise ValueError("No JSON object or array found in input.")

    start = min([pos for pos in [brace_pos, bracket_pos] if pos != -1])
    tail = cleaned[start:]

    def balanced_slice(s):
        stack = []
        for i, c in enumerate(s):
            if c in '{[':
                stack.append(c)
            elif c in '}]':
                if not stack:
                    break
                if ((stack[-1] == '{' and c == '}') or
                    (stack[-1] == '[' and c == ']')):
                    stack.pop()
                else:
                    break
                if not stack:
                    return s[:i+1]
        return s
        
    tail = balanced_slice(tail)

    try:
        return json.loads(tail)
    except json.JSONDecodeError as e:
        raise json.JSONDecodeError(
            f"Failed to parse cleaned JSON. Original: {original[:300]}...\nCleaned: {tail[:300]}...\nError: {e}",
            tail,
            e.pos
        )

@tool
async def save_architecture(content: str, filename: str):
    """
    Function to convert any generated architecture content to a docx file
    MUST BE USED TO SAVE ARCHITECTURES GENERATED CONTENT
    Args:
    content (str): the actual markdown content
    filename (str) : the filename for saving should be of the format project_topic_architecture.docx
    """
    try:
        broadcast_log(manager, filename, level="LOG")
        broadcast_log(manager, "Saving architecture document...", level="INFO")
        loop = asyncio.get_event_loop()
        user_id = get_user_id()
        session_id = get_session_id()
        out_path = f"{esett.FILES}/{user_id}/{get_agent_folder()}/{session_id}/output"
    
        if not os.path.exists(out_path):
            os.makedirs(out_path)
        if not filename:
            filename = "architecture.docx"
        output_path = os.path.join(out_path, filename)
        temp_output_path = output_path.replace(".docx", "temp.docx")
        return_statement = await markdown_to_docx_async(content, temp_output_path)
        
        # Run DocxTemplate operations in executor
        loop = asyncio.get_event_loop()
        
        def _process_template():
            doc = DocxTemplate(temp_output_path)
            print("Mermaid content:", shared.mermaid)
            my_image = InlineImage(doc, image_descriptor=loop.run_until_complete(mermaid_to_png_async(shared.mermaid,session_id)), width=Mm(150))

            context = {"mermaid": my_image}
            print(context,"@########")
            doc.render(context)
            doc.save(output_path)
            return output_path
        
        final_path = await loop.run_in_executor(None, _process_template)
        try:
            session_id = get_session_id()
            print(f"Saving document to {output_path}")
            filename = os.path.basename(output_path)
            await broadcast_file_generated(session_id, filename, output_path)
            print(f"Successfully saved document to {output_path}")

        except Exception as e:
            print(f"An exception occurred while saving the doc file: {e}")
        

        broadcast_log(manager, f"Architecture document saved to {output_path}", level="INFO")
        
        
        return f"Document saved to {output_path}"
        
    except Exception as e:
        error_msg = f"File saved successfully: {str(e)}"
        broadcast_log(manager, error_msg, level="INFO")
        print(error_msg)
        return error_msg

@tool
async def general_query(file_names: List[str], query: str):
    """
    Can be used to generate a response for a general user query. A general user query is one
    that does ask to generate a specific type of document rather is just looking for an answer
   
    Args:
        file_names (List[str]): A list of file names to analyze
                                These names must be obtained from the upload_file tool
        query (str): the users query
    """
    print(f"Executing tool: general_query for files: {file_names}---")
    broadcast_log(manager, "Processing general query...", level="INFO")
    
    gemini_files = []
    for name in file_names:
        if os.path.exists(name):
            async with aiofiles.open(name, "rb") as f:
                file_data = await f.read()
                file = pickle.loads(file_data)
                gemini_files.append(file)
        else:
            return f"Error: file {name} has not yet been uploaded"
            
    if not gemini_files:
        return f"Error: Must provide at least one file to process general query"

    global model
    prompt = query
    try:
        broadcast_log(manager, "Generating response for general query...", level="INFO")
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(None, model.generate_content, [prompt] + gemini_files)
        broadcast_log(manager, "General query processing completed", level="INFO")
        return response.text
    except Exception as e:
        error_msg = f"An exception occurred in general query: {str(e)}"
        broadcast_log(manager, error_msg, level="ERROR")
        return error_msg

# Convert tools to async
tools = [upload_file, delete_file, general_query, update_response, markdowntodoc, generate_architechture, save_architecture]

orchestrator = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=.3, timeout=None, max_retries=2).bind_tools(tools)

def agent(state: AgentState):
    try:
        response = orchestrator.invoke(state["messages"])
        return {"messages": [response]}
    except Exception as e:
        print(f"Exception in agent exception is {e}")
        return "error"

def action(state: AgentState):
    """Synchronous action function that executes async tools"""
    print("Executing tool actions...")
    last_message = state['messages'][-1]

    async def execute_tools():
        results = []
        for tc in last_message.tool_calls:
            broadcast_log(manager, f"Invoking tool {tc['name']} with args {tc['args']}", level="LOGS")
            tool_to_call = next(t for t in tools if t.name == tc['name'])
            # Use ainvoke for async tools
            obs = await tool_to_call.ainvoke(tc['args'])
            results.append(ToolMessage(content=obs, tool_call_id=tc['id'], name=tc['name']))
            print(f"Tool {tc['name']} execution completed")
        return results

    # Handle async execution properly
    try:
        try:
            loop = asyncio.get_running_loop()
            import nest_asyncio
            nest_asyncio.apply()
            tool_invoke_res = asyncio.run(execute_tools())
        except RuntimeError:
            tool_invoke_res = asyncio.run(execute_tools())
    except Exception as e:
        print(f"Error in action function: {e}")
        return {"messages": [ToolMessage(content=f"Error executing tools: {str(e)}", tool_call_id="error", name="error")]}

    print("All tool actions completed")
    return {"messages": tool_invoke_res}

def should_continue(state: AgentState) -> str:
    return "action" if state["messages"][-1].tool_calls else False

# Create workflow
workflow = StateGraph(AgentState)
workflow.add_node("agent", agent)
workflow.add_node("tools", action)
workflow.add_edge("tools", "agent")
workflow.set_entry_point("agent")
workflow.add_conditional_edges("agent", should_continue, {"action": "tools", False: END})

memory = MemorySaver()
app = workflow.compile(checkpointer=memory)

# Async workflow execution functions (matching your working pattern)
async def run_architecture_workflow_async(messages, config=None):
    """Async wrapper that runs the architecture workflow"""
    if config is None:
        config = {"configurable": {"thread_id": "default"}}

    broadcast_log(manager, "Starting async architecture workflow execution...", level="INFO")

    result = None
    async for event in app.astream({"messages": messages}, config):
        broadcast_log(manager, f"Workflow event: {list(event.keys())}", level="DEBUG")
        result = event

    broadcast_log(manager, "Async architecture workflow execution completed", level="INFO")
    return result

def run_architecture_workflow_sync(messages, config=None):
    """Synchronous architecture workflow execution using async API"""
    if config is None:
        config = {"configurable": {"thread_id": "default"}}

    broadcast_log(manager, "Starting architecture workflow execution...", level="INFO")

    # Use asyncio.run to run the async workflow
    result = asyncio.run(run_architecture_workflow_async(messages, config))

    broadcast_log(manager, "Architecture workflow execution completed", level="INFO")
    return result

# Usage examples:
async def main_arch_async():
    """Example async usage"""
    messages = [HumanMessage(content="Generate architecture from the uploaded files")]
    await run_architecture_workflow_async(messages)

def main_arch_sync():
    """Example sync usage"""
    messages = [HumanMessage(content="Generate architecture from the uploaded files")]
    return run_architecture_workflow_sync(messages)
