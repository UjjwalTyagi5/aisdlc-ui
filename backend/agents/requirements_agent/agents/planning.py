"""Requirements & Ingestion Agent â€” LangGraph workflow.

Enterprise-grade agent that walks the user through ingesting Azure DevOps
work items, normalising acceptance criteria to Gherkin format, detecting
requirement gaps, and producing a structured requirements payload for the
Design Agent.

Tools (LLM-callable):
BOARD READ:
  list_board_projects, list_board_groups, list_board_states,
  list_board_items, list_board_items_by_state, fetch_board_item_detail,
  fetch_board_hierarchy
BOARD WRITE:
  commit_board_ingestion_batch, create_board_item, write_stories_to_board,
  write_acceptance_criteria_to_board, write_back_normalized_to_board
FILE:
  read_uploaded_file
GENERATION:
  generate_user_stories, generate_acceptance_criteria,
  normalize_acceptance_criteria, detect_requirement_gaps,
  build_requirements_payload, identify_epics, generate_stories_from_brd

The exported symbol `app` is consumed by `requirements_agent_api.py`
and the orchestrator-side router.
"""
from __future__ import annotations

import asyncio
import os
import re
import sys
import inspect
import uuid
from typing import Annotated, Any, Dict, List, Optional, TypedDict

from langchain_core.messages import AIMessage, BaseMessage, HumanMessage, ToolMessage
from langchain_core.tools import tool
from langchain_anthropic import ChatAnthropic
from langgraph.checkpoint.memory import MemorySaver
from langgraph.graph import END, StateGraph
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode  # noqa: F401  (kept for parity with peer agents)

import anthropic as _anthropic

from config import sdlcSettings
from config.env import AGENTIC_BASE_URL, ANTHROPIC_API_KEY, ANTHROPIC_MODEL  # noqa: F401
from config.ws_helper import broadcast_log, get_session_id, get_user_id, get_provider_kind
from config.connection_manager import manager
from agents.requirements_agent.config import shared
from agents.requirements_agent.config.shared import get_agent_folder

from config.connector_factory import get_connector_for_session
from config.connectors.base import ConnectorNotAvailableError

_anthropic_client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

def _llm_generate(prompt: str) -> str:
    resp = _anthropic_client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=4096,
        temperature=0.3,
        messages=[{"role": "user", "content": prompt}],
    )
    return resp.content[0].text


def _save_markdown_as_docx(content: str, output_path: str) -> None:
    """Convert markdown text to DOCX using python-docx."""
    import re
    from docx import Document as _Document
    doc = _Document()
    for line in content.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", s):
            doc.add_paragraph(re.sub(r"^\d+\.\s*", "", s), style="List Number")
        else:
            p = doc.add_paragraph()
            # bold inline **text**
            parts = re.split(r"\*\*(.+?)\*\*", s)
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
    doc.save(output_path)


# â”€â”€ Connector helper â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Returns (connector, None) on success or (None, error_string) on failure.
# Tools call this instead of reading vendor connector configuration directly.

async def _get_provider(kind: str | None = None):
    resolved = kind or get_provider_kind()
    try:
        connector = await get_connector_for_session(resolved)
        broadcast_log(
            manager,
            f"Using {connector.display_name} board connector ({connector.connector_name})",
            level="INFO",
        )
        return connector, None
    except ConnectorNotAvailableError as exc:
        return None, str(exc)
    except Exception as exc:
        return None, f"Connector error: {exc}"


# â”€â”€ Tools â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

@tool
async def list_board_projects() -> str:
    """List all projects visible in the configured project management connector.

    Returns a human-readable summary of project names. Call this first when
    the user wants to start ingesting stories.
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    try:
        projects = await connector.read_adapter("list_projects")
    except Exception as exc:  # noqa: BLE001
        return f"Error fetching projects: {exc}"
    if not projects:
        return "No projects found."
    lines = [
        f"- {p.get('name', '')}" + (f" ({p.get('key')})" if p.get("key") else "")
        for p in projects
    ]
    return f"Projects available in {connector.display_name}:\n" + "\n".join(lines)


@tool
async def list_board_groups(project: str) -> str:
    """List optional board groups/teams for the given project.

    Jira usually returns no groups here because Jira project roles are
    permissions, not delivery teams. If no groups are returned, continue by
    calling list_board_items(project) without a team filter.

    Args:
        project: Exact project name/key as returned by list_board_projects.
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    try:
        teams = await connector.read_adapter("list_teams", project=project)
    except Exception as exc:  # noqa: BLE001
        return f"Error fetching teams for project '{project}': {exc}"
    if not teams:
        return (
            f"No board groups/teams found in project '{project}'. "
            "Continue without a team filter by calling list_board_items(project)."
        )
    lines = [f"- {t['name']}" for t in teams]
    return f"Teams in '{project}':\n" + "\n".join(lines)


@tool
async def list_board_states(project: str, work_item_type: str = "User Story") -> str:
    """List the workflow states defined for a work item type in the project.

    Use this to discover the real state names (e.g. "New", "Active") so you
    never invent values. Defaults to User Story.
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    try:
        states = await connector.read_adapter(
            "list_states", project=project, item_type=work_item_type
        )
    except Exception as exc:  # noqa: BLE001
        return f"Error fetching states: {exc}"
    if not states:
        return "No states returned."
    lines = [f"- {s['name']} ({s.get('category', '')})" for s in states]
    return f"States for '{work_item_type}' in '{project}':\n" + "\n".join(lines)


@tool
async def list_board_items_by_state(project: str, state: str, team: Optional[str] = None) -> str:
    """List user stories matching a state filter in the project.

    Args:
        project: Project name.
        state: Exact state name from list_board_states (e.g. "New", "Active", "To Do").
        team: Optional team name to scope the query.

    Returns a numbered list with ID, type, and title for each story.
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    try:
        stories = await connector.read_adapter(
            "list_stories", project=project, state=state, team=team
        )
    except Exception as exc:  # noqa: BLE001
        return f"Error fetching stories: {exc}"
    if not stories:
        return f"No stories in state '{state}' for project '{project}'."
    lines = [
        f"#{s['source_key']} (id: {s['id']}) [{s['work_item_type']}] {s['title']}"
        + (f" (assigned: {s['assigned_to']})" if s.get("assigned_to") else "")
        for s in stories
    ]
    return f"Found {len(stories)} stories in '{state}':\n" + "\n".join(lines)


@tool
async def list_board_items(project: str, team: Optional[str] = None) -> str:
    """Fetch ALL work items in a project regardless of state or type.

    Use this FIRST when the user asks to pull existing board items.
    Never guess states â€” this tool returns everything without filters.

    Args:
        project: Project name.
        team: Optional team name to scope the query.
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    try:
        items = await connector.read_adapter(
            "list_all_items", project=project, team=team
        )
    except Exception as exc:
        return f"Error fetching work items: {exc}"
    if not items:
        return f"No work items found in project '{project}'."
    lines = [
        f"#{s['source_key']} (id: {s['id']}) [{s['work_item_type']}] [{s['state']}] {s['title']}"
        + (f" (assigned: {s['assigned_to']})" if s.get("assigned_to") else "")
        for s in items
    ]
    return f"Found {len(items)} work items in '{project}':\n" + "\n".join(lines)


@tool
async def fetch_board_item_detail(project: str, work_item_id: int) -> str:
    """Fetch the full normalized detail of one story (description + acceptance criteria)."""
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    try:
        norm = await connector.read_adapter(
            "fetch_item_detail", project=project, item_id=work_item_id
        )
    except Exception as exc:  # noqa: BLE001
        return f"Error fetching work item #{work_item_id}: {exc}"
    ac_lines = "\n".join(f"  - {c}" for c in norm.get("acceptance_criteria", [])) or "  (none)"
    return (
        f"Story #{norm.get('source_key') or norm['work_item_id']} (id: {norm['id']}): {norm['title']}\n"
        f"State: {norm['state']} | Type: {norm['work_item_type']}\n"
        f"Description:\n{norm.get('description') or '(empty)'}\n"
        f"Acceptance Criteria:\n{ac_lines}"
    )


@tool
async def commit_board_ingestion_batch(
    project: str,
    state_filter: str,
    target_state: str,
    work_item_ids: List[int],
    team: Optional[str] = None,
) -> str:
    """Move selected stories to the target state and persist them as an ingestion batch.

    This is the only mutating tool. Always confirm with the user BEFORE calling it.

    Args:
        project: Project name.
        state_filter: The state stories were filtered by (recorded for audit).
        target_state: The state to move stories to (typically "Active").
        work_item_ids: List of work item IDs the user selected.
        team: Optional team name (recorded for audit).

    Returns batch_id on success.
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    if not work_item_ids:
        return "Error: no work item IDs provided."

    user_id = get_user_id() or "unknown"
    chat_session_id = get_session_id() or "unknown"
    batch_id = f"BATCH-{uuid.uuid4().hex[:12]}"

    org_url = (await connector.auth_adapter()).get("org_url", "")

    moved: List[int] = []
    skipped: List[int] = []
    failed: List[Dict[str, Any]] = []
    normalized_stories: List[Dict[str, Any]] = []

    for wid in work_item_ids:
        try:
            norm = await connector.read_adapter(
                "fetch_item_detail", project=project, item_id=wid
            )
            if norm.get("state") == target_state:
                skipped.append(wid)
            else:
                await connector.write_adapter(
                    "move_item_state", project=project, item_id=wid, new_state=target_state
                )
                norm["state"] = target_state
                moved.append(wid)
            normalized_stories.append(norm)
        except Exception as exc:  # noqa: BLE001
            failed.append({"id": wid, "error": str(exc)})

    if not normalized_stories:
        return f"Error: every story failed. Details: {failed}"

    import json as _json
    summary = [
        f"Batch {batch_id} created with {len(normalized_stories)} stor"
        + ("y" if len(normalized_stories) == 1 else "ies") + ".",
        f"Moved to '{target_state}': {moved or 'none'}",
        f"Already in '{target_state}' (no change): {skipped or 'none'}",
    ]
    if failed:
        summary.append(f"Failed: {failed}")
    summary.append("Ready to hand off to the Design Agent.")
    handoff = {
        "to": "design",
        "batch_id": batch_id,
        "stage_completed": "ingestion",
        "context_keys": ["requirements_payload"],
        "story_count": len(normalized_stories),
        "provider_kind": get_provider_kind(),
    }
    summary.append(f"HANDOFF::{_json.dumps(handoff)}")
    return "\n".join(summary)


@tool
def read_uploaded_file(file_path: str) -> str:
    """Read an uploaded file and return its text content so it can be analysed.

    Supports: .txt, .md, .csv, .xlsx, .xls, .docx, .pdf.
    Call this whenever the user says "use the following files ..." or references a file path.

    Args:
        file_path: Absolute path to the uploaded file as provided in the user message.

    Returns the extracted text content (may be long).
    """
    import pathlib
    import traceback

    path = pathlib.Path(file_path.strip())
    if not path.exists():
        return f"Error: file not found at {file_path}"

    ext = path.suffix.lower()
    try:
        if ext in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="replace")

        if ext == ".csv":
            import csv
            rows = []
            with open(path, newline="", encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f)
                for row in reader:
                    rows.append(", ".join(row))
            return "\n".join(rows)

        if ext in (".xlsx", ".xls"):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
                parts = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    parts.append(f"=== Sheet: {sheet_name} ===")
                    for row in ws.iter_rows(values_only=True):
                        parts.append("\t".join("" if v is None else str(v) for v in row))
                return "\n".join(parts)
            except ImportError:
                import pandas as pd
                df = pd.read_excel(path)
                return df.to_csv(index=False)

        if ext == ".docx":
            from docx import Document
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        if ext == ".pdf":
            try:
                import PyPDF2
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    pages = [reader.pages[i].extract_text() or "" for i in range(len(reader.pages))]
                return "\n".join(pages)
            except ImportError:
                try:
                    import pdfplumber
                    with pdfplumber.open(path) as pdf:
                        return "\n".join(page.extract_text() or "" for page in pdf.pages)
                except ImportError:
                    return f"Error: no PDF library available (install PyPDF2 or pdfplumber)"

        # Fallback: try plain text
        return path.read_text(encoding="utf-8", errors="replace")

    except Exception:
        return f"Error reading file {file_path}:\n{traceback.format_exc()}"


@tool
async def generate_user_stories(stories_text: str, project_context: str = "") -> str:
    """Generate INVEST-formatted user stories from raw story titles/descriptions or BRD text.

    Args:
        stories_text: Raw story list, board item details, or pasted requirements text.
        project_context: Optional extra context (project goal, tech stack, etc.).

    Returns JSON array of user stories with id, title, description, and story_points_estimate.
    """
    prompt = f"""You are a senior business analyst. Given the following input, generate well-formed user stories in INVEST format.

INPUT:
{stories_text}

{"PROJECT CONTEXT: " + project_context if project_context else ""}

For each story output a JSON object with these fields:
- "title": short imperative title (As a <role>, I want <action>)
- "description": 2-3 sentence elaboration of the story
- "story_points_estimate": Fibonacci number (1,2,3,5,8,13) based on complexity
- "invest_check": object with keys Independent, Negotiable, Valuable, Estimable, Small, Testable â€” each true/false
- "notes": any assumptions or gaps identified

Return a JSON array. No markdown, no extra text."""
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, _llm_generate, prompt)
    return result


@tool
async def generate_acceptance_criteria(stories_text: str) -> str:
    """Generate Given/When/Then acceptance criteria for a set of user stories.

    Args:
        stories_text: User stories text or JSON (from generate_user_stories or fetch_board_item_detail).

    Returns JSON array where each item has "story_title" and "acceptance_criteria" (list of Gherkin scenarios).
    """
    prompt = f"""You are a QA lead. For each user story below, write acceptance criteria in Gherkin format (Given/When/Then).

STORIES:
{stories_text}

Output a JSON array. Each element:
- "story_title": the story title
- "acceptance_criteria": list of strings, each a full Gherkin scenario (Given ... When ... Then ...)

No markdown, no extra text. Only valid JSON."""
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, _llm_generate, prompt)
    return result


@tool
async def identify_epics(stories_text: str) -> str:
    """Group user stories into Epics based on theme and business domain.

    Args:
        stories_text: User stories text or JSON.

    Returns JSON array of epics, each with epic_title, epic_description, and list of story_titles.
    """
    prompt = f"""You are a product manager. Group the following user stories into logical Epics.

STORIES:
{stories_text}

Output a JSON array. Each element:
- "epic_title": short name for the epic
- "epic_description": one sentence describing the business goal
- "stories": list of story titles that belong to this epic

No markdown, no extra text. Only valid JSON."""
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, _llm_generate, prompt)
    return result


@tool
async def generate_stories_from_brd(brd_content: str, project_context: str = "") -> str:
    """Generate user stories and acceptance criteria directly from BRD or requirements text.

    Args:
        brd_content: Paste the full BRD / requirements document text here.
        project_context: Optional project name, tech stack, or constraints.

    Returns JSON with "epics" (grouped stories) and "stories" (flat INVEST list with AC).
    """
    prompt = f"""You are a senior business analyst. Analyse the BRD below and extract all functional requirements.
For each requirement produce a user story and acceptance criteria.

BRD:
{brd_content}

{"PROJECT CONTEXT: " + project_context if project_context else ""}

Return a single JSON object:
{{
  "epics": [
    {{"epic_title": "...", "epic_description": "...", "stories": ["story title 1", ...]}}
  ],
  "stories": [
    {{
      "title": "As a <role>, I want <action> so that <benefit>",
      "description": "...",
      "epic": "parent epic title",
      "acceptance_criteria": ["Given ... When ... Then ...", ...],
      "story_points_estimate": <fibonacci number>,
      "invest_check": {{"Independent": true, "Negotiable": true, "Valuable": true, "Estimable": true, "Small": true, "Testable": true}}
    }}
  ]
}}
No markdown. Only valid JSON."""
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, _llm_generate, prompt)
    return result


@tool
async def create_board_item(
    project: str,
    title: str,
    work_item_type: str = "User Story",
    description: str = "",
    acceptance_criteria: str = "",
) -> str:
    """Create a single work item of ANY type in the project management connector.

    Use this to create Epics, Features, User Stories, Tasks, Bugs, etc.

    Args:
        project: Project name.
        title: Work item title.
        work_item_type: Exact type â€” e.g. "Epic", "Feature", "User Story", "Task", "Bug".
        description: Optional description.
        acceptance_criteria: Optional acceptance criteria text.

    Returns the created work item ID and URL.
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    try:
        wi = await connector.write_adapter(
            "create_item",
            project=project,
            item_type=work_item_type,
            title=title,
            description=description,
            acceptance_criteria=acceptance_criteria,
        )
        wid = wi.get("work_item_id") or wi.get("id", "?")
        url = wi.get("url") or wi.get("work_item_url") or wi.get("_links", {}).get("html", {}).get("href", "")
        return f"Created {work_item_type} #{wid}: {title}" + (f"\n{url}" if url else "")
    except Exception as exc:
        return f"Error creating {work_item_type} '{title}': {exc}"


@tool
async def write_stories_to_board(
    stories_json: str,
    project: str,
) -> str:
    """Create generated user stories as new work items in the project management connector.

    Args:
        stories_json: JSON array of stories from generate_user_stories or generate_stories_from_brd.
                      Each story needs at minimum a "title" field.
        project: Project name to create stories in.

    Returns summary of created work item IDs.
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    import json as _json
    try:
        stories = _json.loads(stories_json)
        if isinstance(stories, dict) and "stories" in stories:
            stories = stories["stories"]
    except Exception as exc:
        return f"Error parsing stories JSON: {exc}"

    created, failed = [], []
    for s in stories:
        title = s.get("title", "").strip()
        if not title:
            continue
        description = s.get("description", "")
        ac_lines = s.get("acceptance_criteria", [])
        ac_html = "<br>".join(ac_lines) if isinstance(ac_lines, list) else str(ac_lines)
        try:
            wi = await connector.write_adapter(
                "create_item",
                project=project,
                item_type="User Story",
                title=title,
                description=description,
                acceptance_criteria=ac_html,
            )
            created.append(f"#{wi['id']} {title}")
        except Exception as exc:
            failed.append(f"{title}: {exc}")

    lines = [f"Created {len(created)} work items:"] + created
    if failed:
        lines += [f"\nFailed ({len(failed)}):"] + failed
    return "\n".join(lines)


@tool
async def write_acceptance_criteria_to_board(
    ac_json: str,
    story_ids: List[int],
    project: str,
) -> str:
    """Write generated acceptance criteria back to existing work items.

    Args:
        ac_json: JSON array from generate_acceptance_criteria â€” list of {story_title, acceptance_criteria}.
        story_ids: Ordered list of work item IDs matching the stories in ac_json.
        project: Project name.

    Returns summary of updated work item IDs.
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    import json as _json
    try:
        ac_list = _json.loads(ac_json)
    except Exception as exc:
        return f"Error parsing AC JSON: {exc}"

    updated, failed = [], []
    for idx, wid in enumerate(story_ids):
        if idx >= len(ac_list):
            break
        ac_entry = ac_list[idx]
        criteria = ac_entry.get("acceptance_criteria", [])
        ac_html = "<br>".join(criteria) if isinstance(criteria, list) else str(criteria)
        try:
            await connector.write_adapter(
                "update_item_fields",
                project=project,
                item_id=wid,
                acceptance_criteria=ac_html,
            )
            updated.append(f"#{wid}")
        except Exception as exc:
            failed.append(f"#{wid}: {exc}")

    lines = [f"Updated {len(updated)} work items:"] + updated
    if failed:
        lines += [f"\nFailed ({len(failed)}):"] + failed
    return "\n".join(lines)


@tool
async def fetch_board_hierarchy(project: str) -> str:
    """Fetch the full Epic â†’ Feature â†’ User Story hierarchy from the project management connector.

    Use this when the user wants to see all work items organised by parent-child
    relationships rather than a flat state-filtered list.

    Args:
        project: Exact project name.

    Returns a formatted tree showing Epics, their Features, and each Feature's
    User Stories with state, description, and acceptance criteria counts.
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    try:
        tree = await connector.read_adapter("fetch_hierarchy", project=project)
    except Exception as exc:
        return f"Error fetching hierarchy: {exc}"

    if not tree:
        return f"No work items found in project '{project}'."

    def _render(nodes: list, indent: int = 0) -> list:
        lines = []
        prefix = "  " * indent
        for node in nodes:
            ac_count = len(node.get("acceptance_criteria", []))
            lines.append(
                f"{prefix}[{node['type']}] #{node['id']} {node['title']} "
                f"({node['state']}, {ac_count} AC)"
            )
            if node.get("children"):
                lines.extend(_render(node["children"], indent + 1))
        return lines

    lines = [f"Work item hierarchy for '{project}':"] + _render(tree)
    return "\n".join(lines)


@tool
async def normalize_acceptance_criteria(stories_json: str) -> str:
    """Rewrite raw acceptance criteria from board stories into proper Gherkin Given/When/Then format.

    Call this after fetch_board_item_detail or fetch_board_hierarchy to clean up
    free-form AC text into structured, testable scenarios.

    Args:
        stories_json: JSON array of stories â€” each with "id", "title", and
                      "acceptance_criteria" (list of raw strings from the board provider).

    Returns JSON array with the same stories but acceptance_criteria replaced
    with normalized Gherkin scenarios.
    """
    import json as _json
    prompt = f"""You are a senior QA engineer. Your task is to produce complete, testable acceptance criteria
in Gherkin format (Given/When/Then) for each user story.

INPUT STORIES (JSON):
{stories_json}

Rules:
1. If a story already has acceptance_criteria items: rewrite each into one Gherkin scenario
   ("Given ... When ... Then ..."). Clean up vague phrasing â€” keep the original intent.
2. If a story has an empty acceptance_criteria array: GENERATE 2-4 concrete Gherkin scenarios
   based on the story title and description. Cover the happy path and at least one edge case.
   Do NOT leave any story with an empty acceptance_criteria list.
3. Add an "And" clause when there are multiple outcomes in the same scenario.
4. Do NOT invent requirements beyond what the story title/description implies.
5. Each scenario must be independently testable by a QA engineer.

Return a JSON array. Each element:
{{
  "id": <original work item id>,
  "title": "<story title>",
  "description": "<story description or empty string>",
  "acceptance_criteria": ["Given ... When ... Then ...", ...]
}}

No markdown. Only valid JSON."""
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, _llm_generate, prompt)
    return result


@tool
async def detect_requirement_gaps(stories_json: str) -> str:
    """Analyse a set of stories and detect gaps, ambiguities, and missing non-functional requirements.

    Use this after fetching and normalising stories to ensure requirements are
    complete before handing off to the Design Agent.

    Args:
        stories_json: JSON array of stories with title, description, and
                      acceptance_criteria fields.

    Returns a structured gap report with actionable follow-up questions.
    """
    prompt = f"""You are a senior business analyst reviewing user stories before system design begins.

CRITICAL RULES â€” read before analysing:
1. Only flag a gap if the information is GENUINELY absent from the story text, description, and acceptance criteria.
   If the story already states the answer, do NOT raise a question about it.
2. Do NOT ask about NFRs (performance, security, accessibility, scalability, availability) unless the story
   explicitly involves a novel performance requirement or a new security boundary.
   Standard enterprise portal NFRs (existing role-based access, existing WCAG standards, existing performance
   SLAs) are covered by default â€” do not ask about them.
3. Do NOT ask about scope that is explicitly ruled out in the story (e.g. "Portal only" means don't ask about mobile).
4. Maximum 3 gaps. If you cannot find 3 genuine missing items, return fewer. Do not pad with generic questions.
5. A gap must be something a developer or designer cannot proceed without. If they can make a reasonable
   assumption from the story, it is not a gap.

STORIES:
{stories_json}

Return a JSON object:
{{
  "completeness_score": <0-100, overall quality score>,
  "ready_for_design": <true|false>,
  "gaps": [
    {{
      "story_id": <id or null for global>,
      "story_title": "<title or 'General'>",
      "gap_type": "<missing_ac|ambiguous|scope_gap|conflict>",
      "description": "<what specific information is absent from the story>",
      "follow_up_question": "<exact question to ask â€” only if this cannot be inferred from the story>"
    }}
  ],
  "summary": "<one paragraph overall assessment>"
}}

No markdown. Only valid JSON."""
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, _llm_generate, prompt)
    return result


@tool
async def build_requirements_payload(
    project: str,
    team: str,
    stories_json: str,
    scope_summary: str = "",
    assumptions: str = "",
    out_of_scope: str = "",
) -> str:
    """Build the structured requirements payload that the Design Agent will consume.

    Call this as the FINAL step before emitting the HANDOFF signal. It produces
    a complete, machine-readable summary of all gathered requirements.

    Args:
        project: Board project name/key.
        team: Optional board team/group name.
        stories_json: JSON array of stories with normalised AC (from normalize_acceptance_criteria).
        scope_summary: 1-2 sentence description of what is being built.
        assumptions: Comma-separated list of assumptions (or empty string).
        out_of_scope: Comma-separated list of out-of-scope items (or empty string).

    Returns a JSON payload string to be included in the handoff context.
    """
    import json as _json
    from datetime import datetime as _dt

    try:
        stories = _json.loads(stories_json)
        if isinstance(stories, dict) and "stories" in stories:
            stories = stories["stories"]
    except Exception:
        stories = []

    # Build structured work_items index so downstream agents can update states
    # without regex-parsing the context string and without losing the item type.
    work_items = []
    for s in stories:
        wid = s.get("id") or s.get("work_item_id")
        if wid:
            work_items.append({
                "id": int(wid),
                "source_key": s.get("source_key") or s.get("key") or str(wid),
                "provider_kind": get_provider_kind(),
                "type": s.get("work_item_type", "User Story"),
                "title": s.get("title", ""),
            })

    _kind = get_provider_kind()
    payload = {
        "generated_at": _dt.utcnow().isoformat() + "Z",
        "project": project,
        "team": team,
        "provider_kind": _kind,
        "scope_summary": scope_summary or f"Requirements from project '{project}'",
        "story_count": len(stories),
        "stories": stories,
        "work_items": work_items,
        "assumptions": [a.strip() for a in assumptions.split(",") if a.strip()],
        "out_of_scope": [o.strip() for o in out_of_scope.split(",") if o.strip()],
        "non_functional_requirements": {},
    }

    return f"REQUIREMENTS_PAYLOAD::\n{_json.dumps(payload, indent=2)}"


@tool
async def write_back_normalized_to_board(
    stories_json: str,
    project: str,
    add_audit_comment: bool = True,
) -> str:
    """Write normalised descriptions and Gherkin acceptance criteria back to work items.

    Each story must have an "id" field (work item ID) plus updated
    "description" and/or "acceptance_criteria" fields.
    Optionally adds an audit comment noting the AI normalisation.

    Args:
        stories_json: JSON array of stories from normalize_acceptance_criteria.
        project: Project name.
        add_audit_comment: If True, adds a comment to each updated work item.

    Returns summary of updated work item IDs.
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"

    import json as _json
    from datetime import datetime as _dt

    try:
        stories = _json.loads(stories_json)
        if isinstance(stories, dict) and "stories" in stories:
            stories = stories["stories"]
    except Exception as exc:
        return f"Error parsing stories JSON: {exc}"

    updated, failed = [], []
    audit_note = (
        f"Requirements normalised by SDLC AI Agent on {_dt.utcnow().strftime('%Y-%m-%d %H:%M')} UTC. "
        "Acceptance criteria rewritten to Gherkin (Given/When/Then) format."
    )

    for s in stories:
        wid = s.get("id")
        if not wid:
            continue
        ac_list = s.get("acceptance_criteria", [])
        ac_html = "<br>".join(ac_list) if isinstance(ac_list, list) else str(ac_list)
        description = s.get("description", "")
        try:
            await connector.write_adapter(
                "update_item_fields",
                project=project,
                item_id=wid,
                acceptance_criteria=ac_html,
                description=description,
            )
            if add_audit_comment:
                await connector.write_adapter(
                    "add_comment", project=project, item_id=wid, comment=audit_note
                )
            updated.append(f"#{wid} {s.get('title', '')}")
        except Exception as exc:
            failed.append(f"#{wid}: {exc}")

    lines = [f"Updated {len(updated)} work item(s):"] + updated
    if failed:
        lines += [f"\nFailed ({len(failed)}):"] + failed
    return "\n".join(lines)


@tool
async def move_board_item_state(project: str, work_item_ids: List[int], target_state: str) -> str:
    """Move one or more work items to a specific workflow state without creating a batch.

    Use to update states as work progresses:
      - New â†’ Active when analysis starts
      - Active â†’ Ready for Design when requirements are confirmed

    Always confirm with the user before calling this tool.

    Args:
        project: Project name.
        work_item_ids: List of work item IDs to move.
        target_state: Target state name (use list_board_states to get valid values).
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    moved, failed = [], []
    for wid in work_item_ids:
        try:
            await connector.write_adapter(
                "move_item_state", project=project, item_id=wid, new_state=target_state
            )
            moved.append(f"#{wid}")
        except Exception as exc:
            failed.append(f"#{wid}: {exc}")
    lines = [f"Moved {len(moved)} item(s) to '{target_state}':"] + moved
    if failed:
        lines += [f"Failed ({len(failed)}):"] + failed
    return "\n".join(lines)


@tool
async def add_board_comment(project: str, work_item_id: int, comment: str) -> str:
    """Add a comment to a work item â€” for gap reports, design doc links, PR URLs, or audit notes.

    Args:
        project: Project name.
        work_item_id: Work item ID to comment on.
        comment: Comment text to post.
    """
    connector, err = await _get_provider()
    if err:
        return f"Error: {err}"
    try:
        await connector.write_adapter(
            "add_comment", project=project, item_id=work_item_id, comment=comment
        )
        return f"Comment added to #{work_item_id}."
    except Exception as exc:
        return f"Error adding comment to #{work_item_id}: {exc}"


@tool
async def generate_brd_document(
    project: str,
    scope_summary: str,
    stories_json: str = "",
    gap_report_json: str = "",
) -> str:
    """Generate a Business Requirements Document (BRD) as a downloadable DOCX file.

    Call this after requirements are confirmed and the gap report is reviewed.

    Args:
        project: Project name (used in the document title).
        scope_summary: One or two sentence description of what the project does.
        stories_json: JSON array of stories with Gherkin AC (from normalize_acceptance_criteria).
        gap_report_json: JSON gap report (from detect_requirement_gaps).

    Returns a download URL for the generated DOCX.
    """
    import pathlib as _pathlib
    _stories_block = f"STORIES:\n{stories_json}" if stories_json else ""
    _gap_block = f"GAP REPORT:\n{gap_report_json}" if gap_report_json else ""
    prompt = f"""Write a complete, professional Business Requirements Document for:

PROJECT: {project}
SCOPE: {scope_summary}

{_stories_block}
{_gap_block}

Use this structure:
# Business Requirements Document â€” {project}
## 1. Executive Summary
## 2. Project Overview and Objectives
## 3. Scope
### 3.1 In Scope
### 3.2 Out of Scope
## 4. Stakeholders and Roles
## 5. Functional Requirements
## 6. Non-Functional Requirements
## 7. User Stories Summary
## 8. Acceptance Criteria
## 9. Assumptions and Dependencies
## 10. Open Questions and Risks

Write professional, comprehensive content for each section. Format as markdown."""
    loop = asyncio.get_event_loop()
    content = await loop.run_in_executor(None, _llm_generate, prompt)
    user_id = get_user_id() or "unknown"
    session_id = get_session_id() or "unknown"
    out_dir = _pathlib.Path(__file__).resolve().parents[3] / "files" / user_id / session_id
    out_dir.mkdir(parents=True, exist_ok=True)
    safe = project.replace(" ", "_").replace("/", "-")[:40]
    filename = f"BRD_{safe}.docx"
    _save_markdown_as_docx(content, str(out_dir / filename))
    return f"BRD ready. Download: {AGENTIC_BASE_URL}/generated/{user_id}/{session_id}/{filename}"


@tool
async def generate_user_stories_document(project: str, stories_json: str) -> str:
    """Generate a User Stories document as a downloadable DOCX file.

    Args:
        project: Project name.
        stories_json: JSON array of stories with Gherkin AC (from normalize_acceptance_criteria).

    Returns a download URL for the generated DOCX.
    """
    import pathlib as _pathlib
    prompt = f"""Write a formal User Stories document for project '{project}'.

STORIES DATA:
{stories_json}

Use this structure:
# User Stories â€” {project}
## Overview
## Story Index
## Detailed Stories
(For each: title, description, acceptance criteria in Gherkin Given/When/Then, story points)
## Non-Functional Requirements
## Glossary

Format as markdown."""
    loop = asyncio.get_event_loop()
    content = await loop.run_in_executor(None, _llm_generate, prompt)
    user_id = get_user_id() or "unknown"
    session_id = get_session_id() or "unknown"
    out_dir = _pathlib.Path(__file__).resolve().parents[3] / "files" / user_id / session_id
    out_dir.mkdir(parents=True, exist_ok=True)
    safe = project.replace(" ", "_").replace("/", "-")[:40]
    filename = f"UserStories_{safe}.docx"
    _save_markdown_as_docx(content, str(out_dir / filename))
    return f"User Stories document ready. Download: {AGENTIC_BASE_URL}/generated/{user_id}/{session_id}/{filename}"


@tool
async def generate_risk_register_document(
    project: str,
    stories_json: str = "",
    gap_report_json: str = "",
) -> str:
    """Generate a Risk Register document as a downloadable DOCX file.

    Args:
        project: Project name.
        stories_json: Optional JSON stories for risk context.
        gap_report_json: Optional gap report JSON â€” gaps become risks.

    Returns a download URL for the generated DOCX.
    """
    import pathlib as _pathlib
    _stories_block = f"STORIES CONTEXT:\n{stories_json}" if stories_json else ""
    _gap_block = f"GAP ANALYSIS:\n{gap_report_json}" if gap_report_json else ""
    prompt = f"""Write a professional Risk Register for project '{project}'.

{_stories_block}
{_gap_block}

Use this structure:
# Risk Register â€” {project}
## Overview
## Risk Summary Table
(Risk ID | Description | Likelihood | Impact | Severity | Mitigation | Owner | Status)
## Detailed Risk Entries
(For each: ID, description, likelihood H/M/L, impact H/M/L, severity score,
mitigation strategy, contingency plan, owner, status)
## Risks from Gap Analysis
## Risk Monitoring Plan

Identify risks covering: functional complexity, NFR gaps, integration, security,
timeline, dependencies, and data risks. Format as markdown."""
    loop = asyncio.get_event_loop()
    content = await loop.run_in_executor(None, _llm_generate, prompt)
    user_id = get_user_id() or "unknown"
    session_id = get_session_id() or "unknown"
    out_dir = _pathlib.Path(__file__).resolve().parents[3] / "files" / user_id / session_id
    out_dir.mkdir(parents=True, exist_ok=True)
    safe = project.replace(" ", "_").replace("/", "-")[:40]
    filename = f"RiskRegister_{safe}.docx"
    _save_markdown_as_docx(content, str(out_dir / filename))
    return f"Risk Register ready. Download: {AGENTIC_BASE_URL}/generated/{user_id}/{session_id}/{filename}"


tools = [
    list_board_projects,
    list_board_groups,
    list_board_states,
    list_board_items,
    list_board_items_by_state,
    fetch_board_item_detail,
    fetch_board_hierarchy,
    commit_board_ingestion_batch,
    read_uploaded_file,
    generate_user_stories,
    generate_acceptance_criteria,
    normalize_acceptance_criteria,
    detect_requirement_gaps,
    build_requirements_payload,
    identify_epics,
    generate_stories_from_brd,
    create_board_item,
    write_stories_to_board,
    write_acceptance_criteria_to_board,
    write_back_normalized_to_board,
    move_board_item_state,
    add_board_comment,
    generate_brd_document,
    generate_user_stories_document,
    generate_risk_register_document,
]


# â”€â”€ System prompt â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

INGESTION_SYS_MESSAGE = """\
You are the Requirements & Ingestion Agent. Your sole responsibility is to
gather, normalise, validate, and package project requirements â€” then hand off
to the Design Agent when the user is ready.

â”€â”€ CAPABILITIES â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
A. PM INGESTION        â€” pull stories, features, and epics from {PM_PROVIDER}
B. HIERARCHY VIEW      â€” show full Epic â†’ Feature â†’ User Story tree
C. STORY GENERATION    â€” generate INVEST-formatted user stories from any input
D. AC NORMALISATION    â€” rewrite raw AC into Gherkin Given/When/Then format
E. GAP DETECTION       â€” identify missing AC, NFRs, and ambiguities
F. REQUIREMENTS PAYLOAD â€” build structured JSON for the Design Agent
G. EPIC IDENTIFICATION  â€” group stories into Epics by business theme
H. BRD â†’ STORIES       â€” analyse a pasted BRD and produce full story breakdown
I. WRITE TO PM TOOL     â€” create/update work items and write back normalised content

â”€â”€ ACCURACY RULES (READ FIRST â€” CRITICAL FOR DEMO) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
These rules prevent hallucination. Violating them produces wrong output.

1. ONLY OUTPUT WHAT IS IN THE SOURCE MATERIAL.
   - User stories, AC, business rules, personas: derive ONLY from what the user
     typed, pasted, uploaded, or what the project management tool returned. Never invent them.
   - If a detail (e.g. a rule, a threshold, a role) is not in the source,
     list it in the gap report â€” do NOT silently fill it in.

2. GAP REPORT = HONEST GAPS, NOT INVENTED ANSWERS.
   - detect_requirement_gaps must report actual missing information.
   - Never pre-populate a gap answer yourself. Ask the user. Record their answer.
   - Only call build_requirements_payload AFTER the user has answered every
     follow-up question from the gap report.

3. CONFIRM INFERENCES BEFORE COMMITTING.
   - If you infer something not explicitly stated (e.g. "I assume this is a
     web app"), state the inference and ask: "Is that correct?"
   - Do not proceed until the user confirms or corrects it.

4. GENERATED STORIES MUST TRACE TO INPUT.
   - Every generated user story must cite its source: the BRD section, the
     user's message, or the work item ID it came from.
   - Format: "Source: [Board item #1234 or KEY-123 / BRD Section 3.2 / User provided]"

â”€â”€ SCOPE BOUNDARY (CRITICAL) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
You are responsible for requirements. You do NOT produce design or code yourself.
- If the user asks to proceed to development or build the code: package requirements
  and emit HANDOFF to development. Never block this request.
- If the user asks about HLD, LLD, API design, C4 diagrams, or wireframes:
  say "That's handled by the Design Agent" and emit HANDOFF to design.
- Never produce design documents or code yourself.

â”€â”€ FLOW A: PM INGESTION (Standard) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
1. Call list_board_projects. If one project, use it; otherwise ask user.
   Project output may include both display name and key, e.g. "SDLC Team (SCRUM)".
   Use the exact display name or key returned by the tool.
2. If the provider is Jira, DO NOT require or invent a team. Jira project roles
   like Member, Viewer, Administrator, or atlassian-addons-project-access are
   permissions, not delivery teams. Skip team filtering unless the user gave a
   real board team field explicitly.
3. For Azure DevOps only, call list_board_groups for the project. If one group/team exists, use it; otherwise ask.
4. Call list_board_items(project, team) â€” fetches ALL items regardless
   of state or type. NEVER guess states. This is the correct first call.
   Present results as a numbered list with ID, type, state, and title.
5. Ask which stories to process (user can say "all", "user stories only", or pick IDs).
6. Run FLOW G â€” start with the PRE-STEP to fetch full details, then normalise + gap check.
   Do NOT skip the PRE-STEP even if the user says "process all" â€” you still need to
   fetch each story's description and AC before normalising.
7. When user is satisfied, call commit_board_ingestion_batch to lock in the batch.

RULE: Do NOT call list_board_states + list_board_items_by_state to find initial board items.
      Always use list_board_items â€” it never misses items due to
      unknown state names or unsupported work item types.
RULE: Never print bracketed fake tool narration like "[Called list_board_items]".
      Only summarize results that came from an actual tool response.
RULE: Never write simulated user replies such as "User: Process all stories",
      "User: Member", or any other fake next message. End with a direct question
      to the real user and wait.

â”€â”€ FLOW B: HIERARCHY VIEW â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
- When the user says "show hierarchy", "show epics", "show all work items":
  Call fetch_board_hierarchy(project) to display the provider's parent-child tree.
- Present the tree clearly. Offer to fetch full details for any node.

â”€â”€ FLOW G: NORMALISE + GAP CHECK (run after any story fetch) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

FAST-PATH FOR EXPLICIT DEVELOPMENT INTENT:
If the current user message already says to start/proceed with development,
skip all review prompts and blocking gap-question stops. Fetch full story
details if needed, normalize acceptance criteria, call build_requirements_payload,
then emit HANDOFF to development immediately. Do not ask "does this look
correct?", do not ask gap questions, and do not ask for another confirmation.
You may include assumptions/gaps as non-blocking risks inside the payload/summary.

PRE-STEP â€” FETCH FULL STORY DETAILS (mandatory):
Before normalising, you need description + existing AC for each story.
- The flat list from list_board_items contains only titles â€” it has no description or AC.
- For each User Story (NOT Epics or Features) in the selected set:
  Call fetch_board_item_detail(project, work_item_id) to retrieve the full content.
- Build a JSON array from the results: [{{"id": N, "title": "...", "description": "...", "acceptance_criteria": [...]}}, ...]
  Use this enriched JSON for all subsequent steps. Never run normalize on title-only data.

1. Call normalize_acceptance_criteria on the enriched stories JSON.
   The tool generates Gherkin AC from the story title/description when no AC exists yet.

   Present the results as formatted markdown â€” NOT raw JSON. Use this structure for each story:

   ### Story #[id]: [title]
   **Acceptance Criteria (Gherkin):**
   - Given [context] When [action] Then [outcome]
   - Given [context] When [action] Then [outcome]

   *** STOP HERE. Ask: "Does this look correct? Any changes needed before I check for gaps?" ***
   Wait for the user's answer before proceeding.

2. Call detect_requirement_gaps on the enriched stories JSON.
   Present the completeness score and ONLY the follow-up questions:

   **Completeness Score: [X]/100**
   **Gaps found: [N]**
   1. [Story #id] [question]

   *** STOP HERE. Do NOT call any more tools. Do NOT answer your own questions.
       Your response must end with the questions. Wait for the user's answers. ***

3. After the user answers each question, incorporate their answer.
   If the user says "it's already in the story", "that's defined", "already covered",
   or any equivalent â€” accept it immediately, mark that gap closed, and move to the next.
   Do NOT re-ask or ask for clarification on a closed gap.
   If there are more unanswered questions, ask the next one and STOP again.
   NEVER bundle all questions into one message. One question at a time.

4. Once ALL questions are answered (or dismissed by the user), show the user a summary of the
   finalised requirements and ask: "Requirements look complete â€” proceed to Design, or skip
   straight to Development?"  *** STOP HERE â€” wait for the user's routing answer only. ***

5. As soon as the user answers (e.g. "design" / "development" / "yes, development"):
   call build_requirements_payload immediately, then emit the matching HANDOFF.
   Do NOT ask for a second confirmation. See HANDOFF RULES for the exact sentinel.
6. Offer to write the normalised AC back to the board provider (write_back_normalized_to_board) BEFORE
   emitting the HANDOFF if the user has not already declined it, but do not block the
   HANDOFF on this â€” if they don't respond or say skip, emit HANDOFF immediately.
7. Emit the HANDOFF signal as the very last line of your response.

â”€â”€ FLOW C/D: GENERATE STORIES, AC & EPICS â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
- If user provides story titles/details: call generate_user_stories.
- Call normalize_acceptance_criteria on the result.
- Call detect_requirement_gaps to check quality.
- Call identify_epics to group them.
- Present ALL generated stories, their Gherkin AC, and the gap report to the user.
  *** STOP HERE. Ask the user if the stories look correct before proceeding. ***
- Ask if user wants to push to the board provider via write_stories_to_board.
- Then run FLOW G steps 2â€“7 (gap questions â†’ confirmation â†’ payload â†’ HANDOFF).

CRITICAL â€” SELF-ANSWERING IS FORBIDDEN:
- NEVER answer your own follow-up questions.
- NEVER pre-fill gap answers based on assumptions.
- If you asked a clarifying question, your response MUST END with that question.
  Do not write anything after it. Do not call any more tools until the user replies.

CRITICAL â€” BULK GAP DISMISSAL:
- If the user says anything like "everything is in the story", "all info is already there",
  "that's final", "no more questions needed", or "it's all defined" â€” treat ALL remaining
  gap questions as closed simultaneously.
  Do NOT ask any more gap questions. Jump straight to the finalised summary and HANDOFF prompt.

CRITICAL â€” GAP QUALITY:
- The gap report may contain up to 3 genuine gaps maximum.
- Only ask about something if a developer literally cannot build it without the answer.
- NFRs covered by existing system standards are NOT gaps.

â”€â”€ FLOW H: BRD â†’ STORIES â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
- Ask user to paste the BRD text (or key sections).
- Call generate_stories_from_brd with the pasted content.
- Run normalize_acceptance_criteria and detect_requirement_gaps on the result.
- Present epics, then stories with Gherkin AC and gap report.
- Ask if user wants to create these as board items.

â”€â”€ FLOW I: WRITE TO PM TOOL â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
- create_board_item: creates ONE board item of ANY type.
  If the user says "create an epic", use work_item_type="Epic".
  If they say "create a feature", use work_item_type="Feature".
- write_stories_to_board: bulk-creates User Stories from generated JSON.
- write_back_normalized_to_board: updates existing board items with normalised
  description + Gherkin AC and adds an audit comment.
- write_acceptance_criteria_to_board: patches AC onto existing board item IDs.
- move_board_item_state: move one or more items to a specific state.
  Use after processing starts (â†’ "Active") and after user confirms (â†’ "Ready for Design").
- add_board_comment: post any text as a comment on a board item.
  Use to attach gap report summaries to the parent epic.
- ALWAYS confirm with the user before calling any mutating tool.

â”€â”€ FLOW J: GENERATE DOCUMENTS â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
After requirements are confirmed and the gap report is reviewed, offer to export:
- generate_brd_document(project, scope_summary, stories_json, gap_report_json)
  â†’ full BRD as DOCX, returns download URL
- generate_user_stories_document(project, stories_json)
  â†’ all stories with Gherkin AC as DOCX, returns download URL
- generate_risk_register_document(project, stories_json, gap_report_json)
  â†’ risk register based on stories and gap analysis, returns download URL
Always present the download URL from the tool result as a clickable link in your response.
Post the gap report summary as a comment on the parent epic/item using add_board_comment.

â”€â”€ FILE UPLOADS â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
- When the user's message says "please use the following files ..." followed by
  one or more file paths, call read_uploaded_file for EACH path first.
- Pass extracted text to generate_stories_from_brd or generate_user_stories.
- Supported formats: .txt, .md, .csv, .xlsx, .xls, .docx, .pdf.
- NEVER say you cannot access files â€” always call read_uploaded_file.

â”€â”€ HANDOFF RULES (CRITICAL) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
ALWAYS call build_requirements_payload before emitting any HANDOFF.

ROUTING â€” determine destination from the user's intent:

  â†’ "to":"development"  when user says anything like: "proceed with development",
    "skip design", "just build it", "go to dev", "start coding", "yes" (after
    being asked Design-or-Development and they chose Development), or any message
    expressing intent to go directly to coding. DO NOT block this â€” respect the
    user's choice immediately.

  â†’ "to":"design"  when user says: "proceed to design", "create design",
    "architecture", "yes" (after being asked Design-or-Development and they chose
    Design), or any message expressing intent to create the design first.

The sequence before any HANDOFF:
  a) Call build_requirements_payload (if not already called this session).
  b) Show a brief summary of finalised requirements.
  c) Ask ONE question: "Requirements look complete â€” proceed to Design, or skip
     straight to Development?"  This is the user's ONLY confirmation step.
  d) As soon as the user answers â€” emit the matching HANDOFF immediately.
     Do NOT ask "are you sure?", do NOT ask them to confirm again, do NOT say
     "shall I proceed?" â€” the answer to step c IS the confirmation.
     The frontend Approve button provides a second safety gate; you must not
     also ask for a typed confirmation.

SHORTCUT â€” if the user's message ALREADY expresses routing intent before step c
  (e.g. "proceed with development", "go to dev", "skip design"), skip step c
  entirely and emit HANDOFF immediately after step b.

HANDOFF signals:

1. To Design (user wants design first):
   HANDOFF::{"to":"design","batch_id":"DESIGN-READY","stage_completed":"requirements","context_keys":["requirements_payload"],"triggered_by":"user_confirmed"}

2. To Development (user explicitly skips design):
   HANDOFF::{"to":"development","batch_id":"DEV-READY","stage_completed":"requirements","context_keys":["requirements_payload"],"triggered_by":"user_confirmed"}

3. After commit_board_ingestion_batch succeeds:
   Include the exact HANDOFF:: line returned by that tool. Do not paraphrase.

NEVER say "that's outside my scope" when a user asks to proceed to development.
NEVER tell the user the pipeline order is mandatory â€” the user decides.
The HANDOFF:: line must appear on its own line at the very end of your response.
Do NOT add any text after it. The orchestrator strips it before showing to user.

â”€â”€ RULES â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
- NEVER invent board project/team/state names â€” always use what tools return.
- NEVER claim a tool was called unless the tool actually returned a result in
  this turn. Do not use placeholder text such as "[Called ...]".
- NEVER include fake user messages in your answer. Do not write lines starting
  with "User:" unless quoting actual prior conversation context.
- ALWAYS confirm before any board mutation.
- Run normalize_acceptance_criteria before build_requirements_payload. Run
  detect_requirement_gaps when the user asks for requirements review or has not
  already given explicit development intent. For explicit "start development"
  requests, gaps are non-blocking risks; do not stop the handoff on gap review.
- Keep responses concise and well-formatted. Use bullet lists and headers.
- Present the completeness score and gap count prominently after gap detection.
"""


# â”€â”€ LangGraph workflow â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

class AgentState(TypedDict):
    messages: Annotated[List[BaseMessage], add_messages]


orchestrator = ChatAnthropic(
    model=ANTHROPIC_MODEL,
    api_key=ANTHROPIC_API_KEY,
    temperature=0.2,
    max_tokens=4096,
    max_retries=2,
    extra_headers={"anthropic-beta": "prompt-caching-2024-07-31"},
).bind_tools(tools)


_BOARD_FETCH_RE = re.compile(
    r"\b(fetch|list|pull|show|get)\b.*\b(jira|board|boards|work\s*items?|issues?|stories?|tasks?)\b",
    re.IGNORECASE,
)


def _needs_forced_board_discovery(messages: List[BaseMessage], response: AIMessage) -> bool:
    if getattr(response, "tool_calls", None):
        return False

    last_human_idx = None
    last_human_text = ""
    for idx in range(len(messages) - 1, -1, -1):
        msg = messages[idx]
        if isinstance(msg, HumanMessage):
            last_human_idx = idx
            last_human_text = str(msg.content or "")
            break
    if last_human_idx is None or not _BOARD_FETCH_RE.search(last_human_text):
        return False

    return not any(isinstance(msg, ToolMessage) for msg in messages[last_human_idx + 1 :])


def agent(state: AgentState):
    try:
        response = orchestrator.invoke(state["messages"])
        if _needs_forced_board_discovery(state["messages"], response):
            return {
                "messages": [
                    AIMessage(
                        content="",
                        tool_calls=[
                            {
                                "name": "list_board_projects",
                                "args": {},
                                "id": f"forced_list_projects_{uuid.uuid4().hex[:8]}",
                            }
                        ],
                    )
                ]
            }
        return {"messages": [response]}
    except Exception as e:  # noqa: BLE001
        print(f"Exception in agent: {e}")
        return {"messages": [AIMessage(content=f"Agent error: {e}")]}


def action(state: AgentState):
    """Execute tool calls produced by the LLM."""
    last_message = state["messages"][-1]

    async def execute_tools():
        results = []
        for tc in last_message.tool_calls:
            broadcast_log(manager, f"Invoking tool {tc['name']} with args {tc['args']}", level="LOGS")
            tool_to_call = next(t for t in tools if t.name == tc["name"])
            obs = await tool_to_call.ainvoke(tc["args"])
            results.append(ToolMessage(content=obs, tool_call_id=tc["id"], name=tc["name"]))
        return results

    try:
        try:
            asyncio.get_running_loop()
            import nest_asyncio
            nest_asyncio.apply()
            tool_invoke_res = asyncio.run(execute_tools())
        except RuntimeError:
            tool_invoke_res = asyncio.run(execute_tools())
    except Exception as e:  # noqa: BLE001
        print(f"Error in action: {e}")
        return {"messages": [ToolMessage(content=f"Tool error: {e}", tool_call_id="err", name="error")]}

    return {"messages": tool_invoke_res}


def should_continue(state: AgentState) -> str:
    last = state["messages"][-1]
    return "action" if getattr(last, "tool_calls", None) else False


workflow = StateGraph(AgentState)
workflow.add_node("agent", agent)
workflow.add_node("tools", action)
workflow.add_edge("tools", "agent")
workflow.set_entry_point("agent")
workflow.add_conditional_edges("agent", should_continue, {"action": "tools", False: END})

memory = MemorySaver()
app = workflow.compile(checkpointer=memory)


# â”€â”€ Backward-compatible wrappers (kept for callers that imported these) â”€

async def run_workflow_async(messages, config=None):
    if config is None:
        config = {"configurable": {"thread_id": "default"}}
    result = None
    async for event in app.astream({"messages": messages}, config):
        result = event
    return result


def run_workflow_sync(messages, config=None):
    return asyncio.run(run_workflow_async(messages, config))
