"""Shared markdown → .docx conversion utility.

Core HTML-parsing and docx-building logic lives here.
Agent-specific concerns (mermaid rendering, image fetching, session URL)
are injected via optional async callbacks so this module has no agent imports.

Usage in design agent:
    from shared.tools.docx_tools import markdown_to_docx
    await markdown_to_docx(
        markdown_string, docx_path,
        fetch_image=_fetch_image_bytes,
        render_mermaid=_render_mermaid_to_png,
    )
"""
from __future__ import annotations

import asyncio
import os
import textwrap
from io import BytesIO
from typing import Awaitable, Callable, Optional

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


def _add_hyperlink(paragraph, text: str, url: str):
    """Insert a clickable hyperlink into a docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _inline(paragraph, element, add_hyperlink_fn=_add_hyperlink):
    for child in element.children:
        if isinstance(child, NavigableString):
            if child.strip():
                run = paragraph.add_run(child)
                curr = child.parent
                while curr and curr.name != element.name:
                    if curr.name in ["strong", "b"]:
                        run.bold = True
                    if curr.name in ["em", "i"]:
                        run.italic = True
                    if curr.name == "code":
                        run.font.name = "Courier New"
                    curr = curr.parent
        elif isinstance(child, Tag):
            if child.name == "br":
                paragraph.add_run().add_break()
            elif child.name == "a":
                add_hyperlink_fn(paragraph, child.get_text(), child.get("href", ""))
            else:
                _inline(paragraph, child, add_hyperlink_fn)


async def markdown_to_docx(
    markdown_string: str,
    docx_path: str,
    fetch_image: Optional[Callable[[str], Awaitable[Optional[bytes]]]] = None,
    render_mermaid: Optional[Callable[[str], Optional[bytes]]] = None,
) -> str:
    """Convert a markdown string to a .docx file at docx_path.

    Args:
        markdown_string: Source markdown content.
        docx_path: Absolute output path for the .docx file.
        fetch_image: Optional async callable (url) -> bytes | None for embedding images.
        render_mermaid: Optional sync callable (mermaid_code) -> bytes | None for diagrams.

    Returns a confirmation string.
    """
    dedented = textwrap.dedent(markdown_string).strip()
    doc = Document()
    html = markdown.markdown(dedented, extensions=["extra", "sane_lists", "tables", "fenced_code"])
    soup = BeautifulSoup(html, "html.parser")

    async def _parse(element, doc, list_info=None):
        if not hasattr(element, "name") or element.name is None:
            return

        if element.name in [f"h{i}" for i in range(1, 7)]:
            p = doc.add_heading(level=int(element.name[1]))
            _inline(p, element)

        elif element.name == "p":
            img_child = element.find("img")
            non_img_text = element.get_text().strip()
            if img_child and not non_img_text and fetch_image:
                src = img_child.get("src", "")
                alt = img_child.get("alt", "diagram")
                if src:
                    png_bytes = await fetch_image(src)
                    if png_bytes:
                        try:
                            doc.add_picture(BytesIO(png_bytes), width=Inches(5.5))
                            caption = doc.add_paragraph(f"Figure: {alt}", style="No Spacing")
                            caption.runs[0].italic = True
                            caption.runs[0].font.size = Pt(9)
                        except Exception:
                            doc.add_paragraph(f"[Image: {src}]")
                    else:
                        doc.add_paragraph(f"[Image could not be loaded: {src}]")
            else:
                p = doc.add_paragraph()
                _inline(p, element)

        elif element.name in ["ul", "ol"]:
            style = "List Bullet" if element.name == "ul" else "List Number"
            parent_level = list_info[1] if list_info else -1
            for li in element.find_all("li", recursive=False):
                await _parse(li, doc, (style, parent_level + 1))

        elif element.name == "li" and list_info:
            style, level = list_info
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Inches(0.5 * level)
            _inline(p, element)
            for nested in element.find_all(["ul", "ol"], recursive=False):
                await _parse(nested, doc, list_info)

        elif element.name == "table":
            rows_data = element.find_all("tr")
            if rows_data:
                num_cols = len(rows_data[0].find_all(["th", "td"]))
                table = doc.add_table(rows=len(rows_data), cols=num_cols)
                table.style = "Table Grid"
                for i, row_el in enumerate(rows_data):
                    for j, cell_el in enumerate(row_el.find_all(["th", "td"])):
                        _inline(table.cell(i, j).paragraphs[0], cell_el)

        elif element.name == "pre":
            code_child = element.find("code")
            is_mermaid = (
                code_child is not None
                and "language-mermaid" in (code_child.get("class") or [])
            )
            if is_mermaid and render_mermaid:
                mermaid_code = (code_child or element).get_text()
                loop = asyncio.get_event_loop()
                png_bytes = await loop.run_in_executor(None, render_mermaid, mermaid_code)
                if png_bytes:
                    try:
                        doc.add_picture(BytesIO(png_bytes), width=Inches(5.5))
                        caption = doc.add_paragraph("Figure: Architecture Diagram", style="No Spacing")
                        caption.runs[0].italic = True
                        caption.runs[0].font.size = Pt(9)
                    except Exception:
                        p = doc.add_paragraph(style="No Spacing")
                        run = p.add_run(mermaid_code)
                        run.font.name = "Courier New"
                        run.font.size = Pt(10)
                else:
                    p = doc.add_paragraph(style="No Spacing")
                    run = p.add_run(mermaid_code)
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
            else:
                p = doc.add_paragraph(style="No Spacing")
                run = p.add_run(element.get_text())
                run.font.name = "Courier New"
                run.font.size = Pt(10)

        elif element.name == "hr":
            doc.add_page_break()

    for element in soup.contents:
        await _parse(element, doc)

    os.makedirs(os.path.dirname(os.path.abspath(docx_path)), exist_ok=True)
    doc.save(docx_path)
    return f"Successfully saved document to '{docx_path}'"
