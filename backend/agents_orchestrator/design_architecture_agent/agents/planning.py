import os
import asyncio
import time
import requests
import json
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_COLOR_INDEX
from io import BytesIO
from xhtml2pdf import pisa
import textwrap
import docx
import markdown
from dotenv import load_dotenv
from typing import TypedDict, Annotated, List, Dict, Any
import google.generativeai as genai
from google.ai.generativelanguage import File
from langchain_core.messages import BaseMessage, HumanMessage, ToolMessage, AIMessage
from config.checkpoint import build_checkpointer
from langchain_core.tools import tool
from langgraph.graph import StateGraph, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode
from langgraph.prebuilt import InjectedState
from langgraph.types import Command
from agents_orchestrator.design_architecture_agent.prompts.brd_prompt import BRDPROMPT
from agents_orchestrator.design_architecture_agent.prompts.Pdd_prompt import PDDPROMPT
from agents_orchestrator.design_architecture_agent.prompts.MOM_prompt import MoMPROMPT
from agents_orchestrator.design_architecture_agent.prompts.risk_register_prompt import RISKPROMPT
from docxtpl import DocxTemplate, InlineImage, RichText
import pickle
from agents_orchestrator.design_architecture_agent.config import shared
 
# from main import output_doc

load_dotenv()
genai.configure(api_key=os.environ['GOOGLE_API_KEY'])
config = genai.types.GenerationConfig(temperature=.15)
model = genai.GenerativeModel(model_name="gemini-2.5-flash", )
print("manas is here guys")
def add_hyperlink(paragraph, text, url):
    """
    Adds a hyperlink to a paragraph.

    Args:
        paragraph: The paragraph to add the hyperlink to.
        text (str): The text to display for the link.
        url (str): The URL the link should point to.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    run.append(rPr)
    run.append(OxmlElement('w:t', text=text))
    
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def markdown_to_docx(markdown_string: str, docx_path: str):
    """
    Converts a Markdown or HTML string to a DOCX file with near-full feature support.

    This function is safe for indented multiline strings and handles:
    - All previous features (Headings, Lists, Bold, Italic, Images, Tables, Page Breaks)
    - Hyperlinks
    - Blockquotes
    - Inline code and Fenced Code Blocks
    - Strikethrough text
    """
    # Import OXML elements here to keep the function self-contained
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn

    dedented_markdown = textwrap.dedent(markdown_string).strip()
    doc = Document()
    # Enable all necessary extensions for full feature support
    html = markdown.markdown(dedented_markdown, extensions=['extra', 'sane_lists', 'tables', 'fenced_code', 'codehilite'])
    # with open(docx_path.replace(".docx", ".pdf"), "wb+") as f:  #to save pdf 
    #     pisa.CreatePDF(html, dest=f)
    soup = BeautifulSoup(html, 'html.parser')

    def _add_inline_content(paragraph, element: Tag):
        """Recursively adds formatted runs to a paragraph."""
        for child in element.children:
            if isinstance(child, NavigableString):
                if child.strip():
                    run = paragraph.add_run(child)
                    curr = child.parent
                    # Apply formatting from parent tags
                    while curr is not None and curr.name != element.name:
                        if curr.name in ['strong', 'b']: run.bold = True
                        if curr.name in ['em', 'i']: run.italic = True
                        if curr.name in ['s', 'del']: run.strike = True
                        if curr.name == 'code':
                            run.font.name = 'Courier New'
                        curr = curr.parent
            elif isinstance(child, Tag):
                if child.name == 'br':
                    paragraph.add_run().add_break()
                elif child.name == 'a':
                    add_hyperlink(paragraph, child.get_text(), child.get('href'))
                else:
                    _add_inline_content(paragraph, child)

    def _parse_element(element, doc, list_info=None):
        """Recursively parses a BeautifulSoup element and adds it to the document."""
        if not hasattr(element, 'name') or element.name is None: return

        if element.name in [f'h{i}' for i in range(1, 7)]:
            level = int(element.name[1])
            p = doc.add_heading(level=level)
            _add_inline_content(p, element)
        elif element.name == 'p':
            p = doc.add_paragraph()
            _add_inline_content(p, element)
        elif element.name in ['ul', 'ol']:
            list_style = 'List Bullet' if element.name == 'ul' else 'List Number'
            parent_level = list_info[1] if list_info else -1
            for li in element.find_all('li', recursive=False):
                _parse_element(li, doc, list_info=(list_style, parent_level + 1))
        elif element.name == 'li':
            if list_info:
                style, level = list_info
                p = doc.add_paragraph(style=style)
                p.paragraph_format.left_indent = Inches(0.5 * level)
                _add_inline_content(p, element)
                for nested_list in element.find_all(['ul', 'ol'], recursive=False):
                    _parse_element(nested_list, doc, list_info)
        elif element.name == 'table':
            rows_data = element.find_all('tr')
            if not rows_data: return
            num_cols = len(rows_data[0].find_all(['th', 'td']))
            table = doc.add_table(rows=len(rows_data), cols=num_cols)
            table.style = 'Table Grid'
            for i, row_element in enumerate(rows_data):
                cells = row_element.find_all(['th', 'td'])
                for j, cell_element in enumerate(cells):
                    p = table.cell(i, j).paragraphs[0]
                    _add_inline_content(p, cell_element)
        elif element.name == 'img':
            src = element.get('src')
            if src:
                try:
                    response = requests.get(src, stream=True)
                    response.raise_for_status()
                    doc.add_picture(BytesIO(response.content), width=Inches(5.0))
                except requests.exceptions.RequestException as e:
                    print(f"Warning: Could not fetch image from {src}. Error: {e}")
                    doc.add_paragraph(f"[Image not found: {src}]").add_run().italic = True
        elif element.name == 'hr':
            doc.add_page_break()
        # --- NEW FEATURES ---
        elif element.name == 'blockquote':
            # Add paragraph with "Quote" style and indent
            p = doc.add_paragraph(style='Quote')
            p.paragraph_format.left_indent = Inches(0.5)
            # Blockquotes can contain other elements, so parse children
            for child in element.children:
                _parse_element(child, doc)
        elif element.name == 'pre':
            # Fenced code blocks become <pre><code>...</code></pre>
            code_text = element.get_text()
            p = doc.add_paragraph(style='No Spacing')
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

    # Process all top-level tags from the parsed HTML
    for element in soup.contents:
        _parse_element(element, doc)
    try:
        doc.save(docx_path)
        shared.output_file=docx_path
    except Exception as e:
        print("an exception occured while saving the doc file", e)
    return f"Successfully converted Markdown to '{docx_path}'"

class AgentState(TypedDict):
    messages: Annotated[List[BaseMessage], add_messages]
    
    
@tool

def upload_file(local_path: str):

    """

    Uploads a local file to the cloud service so it can be used by other tools

    It stores the file handle in the agent's state for later use

    Args:

        local_path (str): The local path of the file to upload

    Returns:

        str: The file handle name that can be used by other tools

    """

    print(f"---Executing tool: Upload_file for path: {local_path}---")

    if not os.path.exists(local_path): 

        error_msg = f"Error: local file not found at {local_path}"

        print(error_msg)

        return error_msg

    try:

        print(f"Starting upload for: {local_path}")

        file = genai.upload_file(path=local_path)

        print("File uploaded, waiting for processing...")

        while file.state.name == "PROCESSING":

            time.sleep(2)

            file = genai.get_file(file.name)

            print("Still processing...")

        if file.state.name == "FAILED": 

            error_msg = f"Error: File upload has failed for {local_path}"

            print(error_msg)

            return error_msg

        print(f"File {file.name} is now ACTIVE")

        # Save the file handle
        file_path = "C:/Users/mprasad053/OneDrive - PwC/SDLC/agentic_app/agents/design_architecture_agent"
        final_path = os.path.join(file_path,file.name)
        with open(final_path, "wb+") as f:
            print(f"File handle saved: {final_path}")
            pickle.dump(file, f)

        print(f"File handle saved: {file.name}")

        # Return the file handle name (not a success message)

        return final_path  # ✅ Return the file handle name

    except Exception as e:

        error_msg = f"Exception during upload: {str(e)}"

        print(error_msg)

        return error_msg

@tool
def delete_file(file_name: str):
    """
    Deletes a file from the cloud service using its file name, all uploaded files must be deleted after their use is done
    
    Args:
        file_name (str): The name of the file to delete, use the name returned by upload file
    """    
    print(f"---Executing tool: delete file for name: {file_name}---")
    if not os.path.exists(file_name):
        return f"Error: file {file_name} does not exist"
    genai.delete_file(name=file_name)
    os.remove(file_name)
    print(f"File {file_name} deleted.")
    return f"Succesfully deleted file {file_name}"

@tool 
def generate_brd(file_names: List[str], custom_prompt: str):
    """
    Generates a Business Requirement Document (BRD) by analyzing a collection of files,
    which can include transcripts, audios and meeting recordings.
    
    Args: 
        file_names (List[str]): A list of file names to analyze
                                These names must be obtained from the upload_file tool
        custom_prompt (str): Addtional instructions needed based on users query, can be empty
    """
    print(f"Executing tool: generate_brd for files: {file_names}---")
    gemini_files = []
    for name in file_names:
        if os.path.exists(name):
            with open(name, "rb") as f:
                file = pickle.load(f)
                gemini_files.append(file)
        else:
            return f"Error: file {name} has not yet been uploaded"
    if not gemini_files:
        return f"Error: Must provide atleast one file to generate a BRD"
    #model = genai.GenerativeModel(model_name="gemini-2.5-flash", generation_config=config)
    global model
    prompt = BRDPROMPT.format(custom_prompt=custom_prompt)
    try:
        response = model.generate_content([prompt] + gemini_files) 
        return response.text
        
    except Exception as e:
        print("An exception occured while generating the BRD", e)
        return "An exception occured while generating the BRD"
    #if os.path.exists("outputs/brd.docx"):
    #    os.remove("outputs/brd.docx") 
    #markdown_to_docx(response.text, "outputs/brd.docx")
      
    #return response.text
@tool
def describe_and_upload_av_file(file_path):
    """
    takes an audio or a video file and generates a decription of important details and uploads the file to the cloud service
    Same purpose as upload_file but used for audio and video files.
    Outputs: File handle of uploaded audio or video file.
    Args:
    file_path (str): local path of file to use
    """
    if not os.path.exists(local_path): return f"Error: local file not found at {local_path}"
    file = genai.upload_file(path=local_path)
    while file.state.name == "PROCESSING":
        time.sleep(2)
        file = genai.get_file(file.name)
    if file.state.name == "FAILED": return f"Error: File upload has failed for {local_path}"
    print(f"File {file.name} is now ACTIVE")
    prompt = """Use the attached file to extract important conversational and/or visual details of the meeting"""
    response = model.generate_content([prompt] + file)
    with open(f"summaries/{file_name}.txt", "w+") as f:
        f.write(response.text)
        
    file = upload_file(f"summaries/{file_name}.txt")   
    return file
    
@tool 
def generate_pdd(file_names: List[str], custom_prompt: str):
    """
    Generates a Process Definition Document (PDD) by analyzing a collection of files,
    which can include transcripts, audios and meeting recordings.
    
    Args: 
        file_names (List[str]): A list of file names to analyze
                                These names must be obtained from the upload_file tool
        custom_prompt (str): Additional special instructions to be taken into consideration while generating the PDD
        
    """
    print(f"Executing tool: generate_pdd for files: {file_names}---")
    gemini_files = []
    for name in file_names:
        if os.path.exists(name):
            with open(name, "rb") as f:
                file = pickle.load(f)
                gemini_files.append(file)
        else:
            return f"Error: file {name} has not yet been uploaded"
    if not gemini_files:
        return f"Error: Must provide atleast one file to generate a PDD"
    #model = genai.GenerativeModel(model_name="gemini-2.5-flash")
    global model
    prompt = PDDPROMPT.format(custom_prompt=custom_prompt)
    try:
        response = model.generate_content([prompt] + gemini_files)  
        return response.text
          
    except Exception as e:
        print("An exception occured while generating the PDD", e)
        return "An exception occured while generating the PDD"
    #if os.path.exists("outputs/pdd.docx"):
    #    os.remove("outputs/pdd.docx") 
    #markdown_to_docx(response.text, "outputs/pdd.docx")
    #return response.text

@tool 
def generate_risk_register(file_names: List[str], custom_prompt: str):
    """
    Generates a risk register by analyzing a collection of files,
    which can include transcripts, audios and meeting recordings.
    
    Args: 
        file_names (List[str]): A list of file names to analyze
                                These names must be obtained from the upload_file tool
        custom_prompt (str): Additional special instructions to be taken into consideration while generating the PDD
    """
    print("Executing tool: generate_risk_register for files: {file_names}---")
    gemini_files = []
    for name in file_names:
        if os.path.exists(name):
            with open(name, "rb") as f:
                file = pickle.load(f)
                gemini_files.append(file)
        else:
            return f"Error: file {name} has not yet been uploaded"
    if not gemini_files:
        return f"Error: Must provide atleast one filel to generate a BRD"
    #model = genai.GenerativeModel(model_name="gemini-2.5-flash")
    global model
    prompt = RISKPROMPT.format(custom_prompt=custom_prompt)
    try:
        response = model.generate_content([prompt] + gemini_files)
        return response.text
        
    except Exception as e:
        print("An exception occured while generating the Risk Register", e)
        return "An exception occured while generating the Risk Register"
    #if os.path.exists("outputs/risk_register.docx"):
    #    os.remove("outputs/risk_register.docx") 
    #markdown_to_docx(response.text, "outputs/risk_register.docx")   
    #return response.text
@tool
def template_pdd(content: str, files: List[str]):
    """This is a function that takes a PDD related content and extracts, data from it to be filled into a predefined template
        This function also saves the document it generates
    Args:
        content (str) : The content from which the PDD template fields will be extracted, if file is provided can be empty
        files (List[str]) : A list of file names to analyze These names must be obtained from the upload_file tool 
    """
    gemini_files = []
    for name in files:
        if os.path.exists(name):
            with open(name, "rb") as f:
                file = pickle.load(f)
                gemini_files.append(file)
        else:
            return f"Error: file {name} has not yet been uploaded"
    prompt = f"""
You are a specialist data extraction AI. Your task is to read the unstructured text provided, comprehend its meaning, and populate a precise JSON schema with the information you synthesize.

**CRITICAL DIRECTIVES:**
1.  **ADHERE TO THE SCHEMA:** You MUST use the exact field names as defined in the JSON schema in Part 2. Do not change, add, or omit keys.
2.  **UNSTRUCTURED INPUT:** The source text is free-form (prose, a transcript, etc.). It will likely NOT contain tables. You must infer the data from sentences and paragraphs.
3.  **JSON OUTPUT ONLY:** Your entire response MUST be a single, valid JSON object, without any surrounding text or markdown.
4.  **HANDLE MISSING DATA:** If information for any field or object is not mentioned, use an empty string `""` or an empty list `[]` as appropriate. Do not invent data.

---
**PART 1: CONCEPTUAL ANALYSIS GUIDE**
---
*This section explains how to find the information needed for the specific keys in the schema.*

*   **For lists (`risks`, `assumptions`, `functional_req`, etc.):** Scan the document for mentions of each item. For every distinct item you find (e.g., a specific risk), create a JSON object.
    *   **`"no"`**: If the text numbers the item (e.g., "1. The first risk is..."), use that number. Otherwise, assign your own sequential number (as a string: "1", "2", "3"...).
    *   **`"desc"`**: This is the main description of the item. For a risk, this is the description of the risk itself. For a functional requirement, it's the requirement's description.
    *   **`"impact"`, `"priority"`, `"dependency"`, `"remarks"`**: Find any text that describes these specific attributes for the item and place it in the corresponding field.
*   **For document tables (`invoice`, `seal`, `booking`, `shipping`):** The text will describe these fields in prose. For example, "The invoice's field for the shipper should have a header of 'Exporter'." From this, you would create an object: `{{"no": "1", "field": "shipper", "header": "Exporter", "remarks": ""}}`.
*   **For metadata (`project_title`, `author`, etc.):** Identify these single pieces of information from the text.

---
**PART 2: FINAL JSON OUTPUT SCHEMA (STRICT)**
---
*Your entire output must be a single JSON object that strictly follows this schema.*

```json
{{{{
  "project_title": "(string)",
  "author": "(string)",
  "version": "(string)",
  "date": "(string)",
  "version_table": [
    {{{{
      "version": "(string)",
      "date": "(string)",
      "changes": "(string)",
      "author": "(string)",
      "reviewer": "(string)",
      "approver": "(string)",
      "signoff": "(string)"
    }}}}
  ],
  "introduction": "(string)",
  "intended_audience": "(string)",
  "objective": "(string)",
  "solution_objective": "(string)",
  "functional_table": [
    {{{{
      "business": "(string)",
      "description": "(string)"
    }}}}
  ],
  "steps": [
    {{{{
      "no": "(string)",
      "action": "(string)",
      "details": "(string)"
    }}}}
  ],
  "bl_document": [
    {{{{
      "no": "(string)",
      "document": "(string)",
      "file": "(string)"
    }}}}
  ],
  "invoice": [
    {{{{
      "no": "(string)",
      "field": "(string)",
      "header": "(string)",
      "remarks": "(string)"
    }}}}
  ],
  "seal": [
    {{{{
      "no": "(string)",
      "field": "(string)",
      "header": "(string)",
      "remarks": "(string)"
    }}}}
  ],
  "booking": [
    {{{{
      "no": "(string)",
      "field": "(string)",
      "header": "(string)",
      "remarks": "(string)"
    }}}}
  ],
  "shipping": [
    {{{{
      "no": "(string)",
      "field": "(string)",
      "header": "(string)",
      "remarks": "(string)"
    }}}}
  ],
  "business_logic": [
    {{{{
      "no": "(string)",
      "field": "(string)",
      "business_rule": "(string)",
      "remarks": "(string)"
    }}}}
  ],
  "success_criteria": [
    {{{{
      "criteria": "(string)",
      "target": "(string)",
      "description": "(string)"
    }}}}
  ],
  "success_scope": "(string)",
  "functional_req": [
    {{{{
      "no": "(string)",
      "desc": "(string)",
      "priority": "(string)",
      "dependency": "(string)"
    }}}}
  ],
  "out_of_scope": [
    {{{{
      "no": "(string)",
      "desc": "(string)",
      "priority": "(string)",
      "dependency": "(string)"
    }}}}
  ],
  "assumptions": [
    {{{{
      "no": "(string)",
      "desc": "(string)",
      "impact": "(string)"
    }}}}
  ],
  "risks": [
    {{{{
      "no": "(string)",
      "desc": "(string)",
      "impact": "(string)"
    }}}}
  ],
  "constraints": [
    {{{{
      "no": "(string)",
      "desc": "(string)"
    }}}}
  ],
  "dependencies": [
    {{{{
      "no": "(string)",
      "desc": "(string)",
      "impact": "(string)"
    }}}}
  ]
}}}}"""
    #model = genai.GenerativeModel(model_name="gemini-2.5-flash")
    global model
    response = model.generate_content([prompt] +gemini_files + [content])
    s = response.text.replace("```json", "")
    s = s.replace("```", "")
    s = json.loads(s)
    file = "PDD_Template.docx"
    doc = DocxTemplate(file)
    doc.render(s)
    word_doc = doc.docx
    output_path = "outputs/pdd.docx"
    doc.save(output_path)
    return f"Document saved to {output_path}"
@tool
def update_response(file_names: List[str], query: str, content: str):
    """
    A function that can be used to update the content provided based on context
    
    Args:
    file_names (List[str]): A list of file names obtained from upload file that is provided as context
    query (str): the actual query or the updates the user needs, combine all the needs previously mentioned
    content (str): the actual content that needs to be updated
    """
    print(f"Executing tool: update_response for files: {file_names}---")
    gemini_files = []
    for name in file_names:
        if os.path.exists(name):
            with open(name, "rb") as f:
                file = pickle.load(f)
                gemini_files.append(file)
        else:
            return f"Error: file {name} has not yet been uploaded"
    #model = genai.GenerativeModel(model_name="gemini-2.5-flash")
    global model
    prompt = """Given the following content and attachted files as context update the content provided to you based on the query
Query: {query}

````````
Content:
{content}
`````````
    """
    prompt = prompt.format(content = content, query = query)
    try:
        response = model.generate_content([prompt] + gemini_files)
        return response.text
        
    except Exception as e:
        print("An exception occured while using update query", e)
        return "An exception occured while updating the query"
    #return response.text
    
@tool 
def generate_mom(file_names: List[str], custom_prompt: str):
    """
    Generates the Minutes of the meeting by analyzing a collection of files,
    which can include transcripts, audios and meeting recordings.
    
    Args: 
        file_names (List[str]): A list of file names to analyze
                                These names must be obtained from the upload_file tool
        custom_prompt (str): Additional special instructions to be taken into consideration while generating the PDD
    """
    print("Executing tool: generate_mom for files: {file_names}---")
    gemini_files = []
    for name in file_names:
        if os.path.exists(name):
            with open(name, "rb") as f:
                file = pickle.load(f)
                gemini_files.append(file)
        else:
            return f"Error: file {name} has not yet been uploaded"
    if not gemini_files:
        return f"Error: Must provide atleast one file to generate the minutes of the meeting"
    #model = genai.GenerativeModel(model_name="gemini-2.5-flash")
    global model
    prompt = MoMPROMPT.format(custom_prompt=custom_prompt)
    try:
        response = model.generate_content([prompt] + gemini_files)  
        return response.text
    except Exception as e:
        print("An exception occured while generating the MOM", e)
        return "An exception occured while generating the MOM"
    #if os.path.exists("outputs/mom.docx"):
    #    os.remove("outputs/mom.docx") 
    #markdown_to_docx(response.text, "outputs/mom.docx")  

@tool 
def general_query(file_names: List[str], query: str):
    """
    Can be used to generate a response for a general user query. A general user query is one
    that does ask to generate a specific type of document rather is just looking for an answer
    
    Args: 
        file_names (List[str]): A list of file names to analyze
                                These names must be obtained from the upload_file tool
        query (str): the users query
    """
    print(f"Executing tool: general_query for files: {file_names}---")
    gemini_files = []
    for name in file_names:
        if os.path.exists(name):
            with open(name, "rb") as f:
                file = pickle.load(f)
                gemini_files.append(file)
        else:
            return f"Error: file {name} has not yet been uploaded"
    if not gemini_files:
        return f"Error: Must provide atleast one filel to generate a BRD"
    #model = genai.GenerativeModel(model_name="gemini-2.5-flash")
    global model
    prompt = query
    try:
        response = model.generate_content([prompt] + gemini_files)    
    except Exception as e:
        print("an exception occured in general query", e)
    return response.text

@tool 
def markdowntodoc(content: str, output_path: str):
    """
    Function to convert a HTMl or Markdown string into a docx file and save it
    Args:
    content (str): the actual markdown string
    output_path (str): the local output path to store the file: typically outputs/file.docx
    """
    return markdown_to_docx(content, output_path)


tools = [upload_file, delete_file, generate_brd, generate_mom, generate_pdd, generate_risk_register, general_query, update_response, markdowntodoc, template_pdd]
#, generate_brd, generate_pdd,
         #generate_mom, generate_risk_register, describe_and_upload_av_file, general_query

# NOTE: This module is a legacy/orphaned copy of the design agent graph. The live
# design serving path is agents_orchestrator/design_architecture_agent/agents/architecture.py
# (mounted via design_router_orchestrator). The previous module-level ChatLiteLLM
# singleton was built at import with the shared platform LITELLM key; it has been
# removed as part of BYOK P3.6 so no platform-key client exists here. If this graph
# is ever revived, build the client per-request from resolve_model_for_run(tenant_id,
# model_id) inside agent() as architecture.py does.

def agent(state: AgentState):
    from shared.services.model_resolver import (
        resolve_model_for_run,
        NoModelConfiguredError,
        ModelNotEnabledError,
    )
    try:
        resolved = asyncio.run(
            resolve_model_for_run(state.get("tenant_id", ""), state.get("model_id"))
        )
    except (NoModelConfiguredError, ModelNotEnabledError) as e:
        return {"messages": [AIMessage(content=f"No usable model is configured for this run: {e}")]}
    # Deferred: importing litellm costs ~7s. sys.modules makes repeat calls free.
    from langchain_litellm import ChatLiteLLM
    from shared.services.model_resolver import litellm_key_kwargs  # noqa: PLC0415
    orchestrator = ChatLiteLLM(
        model=resolved.model,
        custom_llm_provider=resolved.litellm_provider,
        api_base=resolved.base_url,
        api_key=resolved.api_key,
        # The BYOK key must be the one litellm uses — see
        # shared/services/model_resolver.litellm_key_kwargs.
        **litellm_key_kwargs(resolved.litellm_provider, resolved.api_key),
        temperature=0.3,
        max_retries=2,
    ).bind_tools(tools)
    try:
        response = orchestrator.invoke(state["messages"])
        return {"messages": [response]}
    except Exception as e:
        print(f"Exception in agent exception is {e}")
        return "error"
tool_node = ToolNode(tools=tools)

def action(state: AgentState):
    last_message = state['messages'][-1]
    tool_invoke_res = []
    for tc in last_message.tool_calls:
        print(f"Invoking tool {tc['name']} with args {tc['args']}")
        tool_to_call = next(t for t in tools if t.name == tc['name'])
        obs = tool_to_call.invoke(tc['args'])
        tool_invoke_res.append(ToolMessage(content=obs, tool_call_id=tc['id'], name=tc['name']))
        #tool_invoke_res.append(AIMessage(content="done calling this tool, proceeding with rest of the query", tool_calls=[tc]))
    return {"messages" : tool_invoke_res}

def should_continue(state: AgentState) -> str:
    return "action" if state["messages"][-1].tool_calls else False
workflow = StateGraph(AgentState)
workflow.add_node("agent", agent)
workflow.add_node("tools", tool_node)
workflow.add_edge("tools", "agent")
workflow.set_entry_point("agent")
workflow.add_conditional_edges("agent", should_continue, {"action" : "tools", False : END})

#workflow.add_edge("agent", END)
app = workflow.compile(checkpointer=build_checkpointer("design"))